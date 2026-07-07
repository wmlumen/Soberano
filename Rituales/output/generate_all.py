import os, re
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

BASE = os.path.join(os.environ['TEMP'])
OUTPUT = r'C:\Temp\kitian\GITHUT\RItoMemphisMisraim\Rituales\output'

ACCENT_FIXES = {
    'ä': 'a', 'ë': 'e', 'ï': 'i', 'ö': 'o', 'ü': 'u',
    'Ä': 'A', 'Ë': 'E', 'Ï': 'I', 'Ö': 'O', 'Ü': 'U',
}

def fix_accents(text):
    """Fix common OCR accent corruptions"""
    # The OCR output has \ufffd (replacement char) for various accented chars
    text = text.replace('\ufffd', '')
    text = text.replace('\u2192', '\u2794')  # arrow
    return text

def fix_spanish(text):
    """Fix known Spanish character issues"""
    text = text.replace('iSois', '\u00bfSois')
    text = text.replace('iQu', '\u00bfQu')
    text = text.replace('iPor qu', '\u00bfPor qu')
    text = text.replace('iC', '\u00bfC')
    text = text.replace('iLa', '\u00bfLa')
    text = text.replace('iDon', '\u00bfDon')
    text = text.replace('iEn qu', '\u00bfEn qu')
    text = text.replace('iExplicad', '\u00bfExplicad')
    text = text.replace('iDesde', '\u00bfDesde')
    text = text.replace('iComo', '\u00bfC\u00f3mo')
    text = text.replace('iNo', '\u00bfNo')
    text = text.replace('iHab', '\u00bfHab')
    text = text.replace('iA qu', '\u00bfA qu')
    text = text.replace('iSu', '\u00bfSu')
    text = text.replace('iCual', '\u00bfCu\u00e1l')
    text = text.replace('iTen', '\u00bfTen')
    text = text.replace('iTrabaj', '\u00bfTrabaj')
    text = text.replace('iLo prom', '\u00bfLo prom')
    text = text.replace('iLo juram', '\u00bfLo juram')
    text = text.replace('iQue', '\u00bfQue')
    text = text.replace('iQui', '\u00bfQui')
    text = text.replace('iPor ', '\u00bfPor ')
    text = text.replace('iEst', '\u00bfEst')
    text = text.replace('iSobre', '\u00bfSobre')
    text = text.replace('iNo esper', '\u00bfNo esper')
    text = text.replace('iOs sug', '\u00bfOs sug')
    text = text.replace('iPued', '\u00bfPued')
    text = text.replace('iPrest', '\u00bfPrest')
    text = text.replace('iSois', '\u00bfSois')
    
    # Fix ;C -> \u00bfC
    text = re.sub(r';C', '\u00bfC', text)
    text = re.sub(r';Como', '\u00bfC\u00f3mo', text)
    text = re.sub(r';Donde', '\u00bfD\u00f3nde', text)
    
    return text

def fix_text_line(line):
    """Fix a line of text - apply Spanish fixes and common OCR errors"""
    # First fix Spanish characters
    line = fix_spanish(line)
    # Fix known OCR misreads
    line = line.replace('j', '\u00bf') if line.startswith('j') and 'Pregunta' not in line else line
    return line

def read_ritual(name):
    path = os.path.join(BASE, 'pdf_imgs' if name == 1 else 'pdf_imgs_2', f'ritual_{name}_completo.txt')
    with open(path, 'r', encoding='utf-8') as f:
        return f.read()

def clean_text(text):
    text = fix_accents(text)
    lines = text.split('\n')
    cleaned = []
    skip_patterns = [
        r'^RITUALES, CEREMONIAS E INSTRUCCIONES',
        r'^HISTORIA\. GRANDES CONSTITUCIONES',
        r'^GRAN LOGIA DE ESPANA DE MENFIS-MIZRAIM',
        r'^\d+$',
        r'^LONSEO$',
        r'^ONSEO$',
        r'^CONSEIO$',
        r'^CONSEJO$'
    ]
    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue
        skip = False
        for pat in skip_patterns:
            if re.match(pat, stripped):
                skip = True
                break
        if not skip:
            cleaned.append(fix_text_line(stripped))
    return '\n'.join(cleaned)

def parse_ritual(text):
    """Parse ritual text into sections of (type, content)"""
    text = clean_text(text)
    sections = []
    lines = text.split('\n')
    i = 0
    while i < len(lines):
        line = lines[i]
        
        # Detect headers
        if line.startswith('CATECISMO') or line.startswith('INSTRUCCION') or 'CATECISMO' in line:
            sections.append(('header', line))
        elif line.startswith('Pregunta') or line.startswith('\u00bfPregunta'):
            q_lines = [line]
            i += 1
            while i < len(lines) and not lines[i].startswith('Respuesta') and not lines[i].startswith('Pregunta') and '\u00bfPregunta' not in lines[i]:
                if not lines[i].startswith('GRAN LOGIA'):
                    q_lines.append(lines[i])
                i += 1
            sections.append(('question', '\n'.join(q_lines)))
            if i < len(lines) and lines[i].startswith('Respuesta'):
                r_lines = [lines[i]]
                i += 1
                while i < len(lines) and not lines[i].startswith('Pregunta') and '\u00bfPregunta' not in lines[i] and not lines[i].startswith('CATECISMO') and not lines[i].startswith('INSTRUCCION'):
                    if not lines[i].startswith('GRAN LOGIA'):
                        r_lines.append(lines[i])
                    i += 1
                sections.append(('answer', '\n'.join(r_lines)))
            continue
        else:
            sections.append(('text', line))
        i += 1
    return sections

HTML_TEMPLATE = '''<!DOCTYPE html>
<html lang="es">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{title}</title>
    <style>
        * {{ margin: 0; padding: 0; box-sizing: border-box; }}
        body {{ font-family: 'Times New Roman', Times, serif; background: #f5f0e8; color: #1a1a1a; padding: 40px 20px; }}
        .container {{ max-width: 800px; margin: 0 auto; background: #fffdf9; padding: 60px 80px; box-shadow: 0 2px 20px rgba(0,0,0,0.1); border: 1px solid #d4c9b0; }}
        .gloria {{ font-size: 14px; font-variant: small-caps; letter-spacing: 2px; color: #8b7355; margin-bottom: 20px; text-align: center; }}
        .rito {{ font-size: 16px; font-weight: bold; text-transform: uppercase; letter-spacing: 3px; margin-bottom: 5px; text-align: center; }}
        .orden {{ font-size: 14px; font-style: italic; color: #555; margin-bottom: 15px; text-align: center; }}
        .gran-logia {{ font-size: 13px; font-weight: bold; color: #8b7355; margin-bottom: 5px; text-align: center; }}
        .instruccion {{ font-size: 18px; font-weight: bold; text-align: center; margin: 30px 0 20px; }}
        .catecismo {{ font-size: 16px; font-weight: bold; text-align: center; margin: 20px 0; font-variant: small-caps; }}
        .question {{ margin: 15px 0 5px; font-style: italic; color: #2c2c2c; }}
        .question strong {{ color: #8b7355; }}
        .answer {{ margin: 0 0 15px 20px; padding-left: 15px; border-left: 3px solid #d4c9b0; color: #333; }}
        .answer strong {{ color: #8b7355; }}
        .text-block {{ margin: 10px 0; line-height: 1.6; text-align: justify; }}
        .aclamacion {{ text-align: center; font-weight: bold; margin: 15px 0; letter-spacing: 1px; }}
        .stage {{ font-style: italic; color: #666; margin: 10px 0; padding-left: 20px; }}
        .oracion {{ text-align: center; font-style: italic; margin: 15px 0; }}
        .footer {{ text-align: center; margin-top: 40px; padding-top: 20px; border-top: 1px solid #d4c9b0; font-size: 12px; color: #888; }}
        @media print {{ body {{ background: white; padding: 0; }} .container {{ box-shadow: none; border: none; padding: 40px; }} }}
    </style>
</head>
<body>
    <div class="container">
        {content}
    </div>
</body>
</html>'''

def is_stage_direction(line):
    """Check if a line is a stage direction (in parentheses or starts with special marker)"""
    return line.startswith('\u2794') or line.startswith('(*)') or (line.startswith('(') and line.endswith(')'))

def build_html_content(sections):
    html_parts = []
    for stype, scontent in sections:
        if stype == 'header':
            if 'CATECISMO' in scontent:
                html_parts.append(f'<div class="catecismo">{scontent}</div>')
            else:
                html_parts.append(f'<div class="instruccion">{scontent}</div>')
        elif stype == 'question':
            q = scontent.replace('Pregunta', '<strong>Pregunta</strong>')
            html_parts.append(f'<div class="question">{q}</div>')
        elif stype == 'answer':
            a = scontent.replace('Respuesta', '<strong>Respuesta</strong>')
            html_parts.append(f'<div class="answer">{a}</div>')
        elif stype == 'text':
            if 'A la Gloria' in scontent:
                html_parts.append(f'<div class="gloria">{scontent}</div>')
            elif 'RITO ANTIGUO' in scontent:
                # Fix missing spaces
                s = scontent.replace('ANTIGUOYPRIMITIVODEMENFIS', 'ANTIGUO Y PRIMITIVO DE MENFIS')
                html_parts.append(f'<div class="rito">{s}</div>')
            elif 'ORDEN DE LOS' in scontent:
                html_parts.append(f'<div class="orden">{scontent}</div>')
            elif 'GRAN LOGIA' in scontent:
                html_parts.append(f'<div class="gran-logia">{scontent}</div>')
            elif 'Ritos reunificados' in scontent:
                html_parts.append(f'<div style="font-style:italic;text-align:center;margin:10px 0;">{scontent}</div>')
            elif scontent.startswith('CHEQUETET') or scontent.startswith('MAK-BENAH') or scontent.startswith('\u00bfCHEQUETET'):
                html_parts.append(f'<div class="aclamacion">{scontent}</div>')
            elif is_stage_direction(scontent):
                html_parts.append(f'<div class="stage">{scontent}</div>')
            elif scontent.startswith('Am\u00e9n') or scontent.startswith('Soberano Maestro'):
                html_parts.append(f'<div class="oracion">{scontent}</div>')
            else:
                html_parts.append(f'<div class="text-block">{scontent}</div>')
    return '\n'.join(html_parts)

def generate_html(rnumber, title, sections, output_dir):
    content = build_html_content(sections)
    full_title = f"Ritual {rnumber} - {title}"
    html = HTML_TEMPLATE.format(title=full_title, content=content)
    safe_title = title.replace(' ', '_').replace('\u00ba', 'o').replace('/', '_')
    fname = f'Ritual_{rnumber}_{safe_title}.html'
    path = os.path.join(output_dir, fname)
    with open(path, 'w', encoding='utf-8') as f:
        f.write(html)
    print(f'  HTML: {fname}')
    return path

def generate_md(rnumber, title, sections, output_dir):
    md_parts = []
    md_parts.append(f'# {title}')
    md_parts.append('')
    md_parts.append('---')
    md_parts.append('')
    
    for stype, scontent in sections:
        if stype == 'header':
            if 'CATECISMO' in scontent:
                md_parts.append(f'## {scontent}')
            else:
                md_parts.append(f'### {scontent}')
            md_parts.append('')
        elif stype == 'question':
            q = scontent.replace('Pregunta', '**Pregunta:**')
            md_parts.append(q)
            md_parts.append('')
        elif stype == 'answer':
            a = scontent.replace('Respuesta', '**Respuesta:**')
            md_parts.append(f'> {a}')
            md_parts.append('')
        elif stype == 'text':
            if 'A la Gloria' in scontent:
                md_parts.append(f'*{scontent}*')
            elif 'RITO ANTIGUO' in scontent:
                s = scontent.replace('ANTIGUOYPRIMITIVODEMENFIS', 'ANTIGUO Y PRIMITIVO DE MENFIS')
                md_parts.append(f'# {s}')
            elif 'ORDEN DE LOS' in scontent or 'Ritos reunificados' in scontent:
                md_parts.append(f'*{scontent}*')
            elif 'GRAN LOGIA' in scontent:
                md_parts.append(f'**{scontent}**')
            elif scontent.startswith('CHEQUETET') or scontent.startswith('MAK-BENAH'):
                md_parts.append(f'**{scontent}**')
            elif is_stage_direction(scontent):
                md_parts.append(f'*{scontent}*')
            else:
                md_parts.append(scontent)
            md_parts.append('')
    
    safe_title = title.replace(' ', '_').replace('\u00ba', 'o').replace('/', '_')
    fname = f'Ritual_{rnumber}_{safe_title}.md'
    path = os.path.join(output_dir, fname)
    with open(path, 'w', encoding='utf-8') as f:
        f.write('\n'.join(md_parts))
    print(f'  MD:   {fname}')
    return path

def generate_docx(rnumber, title, sections, output_dir):
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    
    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(16)
    p.add_run().add_break()
    
    for stype, scontent in sections:
        if stype == 'header':
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(scontent)
            run.bold = True
            run.font.size = Pt(13)
        elif stype == 'question':
            q = scontent.replace('Pregunta', 'Pregunta:')
            p = doc.add_paragraph()
            run = p.add_run(q)
            run.italic = True
        elif stype == 'answer':
            a = scontent.replace('Respuesta', 'Respuesta:')
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            p.add_run(a)
        elif stype == 'text':
            p = doc.add_paragraph()
            align_center = any(x in scontent for x in ['Gloria', 'RITO', 'ORDEN', 'GRAN LOGIA', 'CHEQUETET', 'MAK-BENAH', 'Ritos', 'Am\u00e9n', 'Soberano'])
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if align_center else WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(scontent)
            if 'A la Gloria' in scontent or 'Ritos reunificados' in scontent or is_stage_direction(scontent):
                run.italic = True
            if 'RITO ANTIGUO' in scontent:
                run.bold = True
            if 'GRAN LOGIA' in scontent:
                run.bold = True
            if scontent.startswith('CHEQUETET') or scontent.startswith('MAK-BENAH'):
                run.bold = True
    
    # Footer
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run('GRAN LOGIA DE ESPA\u00d1A DE MENFIS-MIZRAIM 2011').font.size = Pt(10)
    
    safe_title = title.replace(' ', '_').replace('\u00ba', 'o').replace('/', '_')
    fname = f'Ritual_{rnumber}_{safe_title}.docx'
    path = os.path.join(output_dir, fname)
    doc.save(path)
    print(f'  DOCX: {fname}')
    return path

rituales = [
    (1, 'Instruccion del 1er Grado - Aprendiz'),
    (2, 'Instruccion del 2o Grado - Companero'),
    (3, 'Ritual e Instruccion del 3er Grado - Maestro'),
]

html_dir = os.path.join(OUTPUT, 'html')
md_dir = os.path.join(OUTPUT, 'md')
doc_dir = os.path.join(OUTPUT, 'doc')

print("Generando rituales...")
for rnum, rtitle in rituales:
    print(f'\nRitual {rnum}: {rtitle}')
    text = read_ritual(rnum)
    sections = parse_ritual(text)
    print(f'  Secciones: {len(sections)}')
    generate_html(rnum, rtitle, sections, html_dir)
    generate_md(rnum, rtitle, sections, md_dir)
    generate_docx(rnum, rtitle, sections, doc_dir)

print('\n\u00a1Todos los rituales generados exitosamente!')
