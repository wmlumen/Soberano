import os, re, sys
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT = r'C:\Temp\kitian\GITHUT\RItoMemphisMisraim\Rituales\output'
SOURCE = os.path.join(os.environ['TEMP'], 'pdf_text_new', 'ritual_4_completo.txt')

DOCUMENTS = [
    {
        'id': 'Ritual_1_Instruccion_1er_Grado_Aprendiz',
        'title': 'Instrucción del 1er Grado – Aprendiz',
        'pages': [68, 70, 71, 73, 74, 75],
    },
    {
        'id': 'Ritual_2_Instruccion_2o_Grado_Companero',
        'title': 'Instrucción del 2º Grado – Compañero',
        'pages': [133, 134, 135, 136],
    },
    {
        'id': 'Ritual_3_Maestro',
        'title': 'Ritual y Instrucción del 3er Grado – Maestro',
        'pages': list(range(137, 178)),
    },
    {
        'id': 'Grandes_Constituciones',
        'title': 'Grandes Constituciones y Reglamentos Generales',
        'pages': list(range(321, 391)),
    },
    {
        'id': 'Rituales_Simbolicos_Completo',
        'title': 'Rituales e Instrucciones de los Grados Simbólicos – Obra Completa',
        'pages': list(range(2, 391)),
    },
]

def load_source():
    with open(SOURCE, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Reemplazar el nombre de la Gran Logia usando regex para manejar saltos de línea
    import re
    
    # Patrón principal: maneja saltos de línea y espacios variables
    pattern = r'Gran\s+Logia\s+de\s+Espa[ñn]a\s+de\s+Menfis[-\s]+Mizraim\s+del\s+Rito\s+Antiguo\s+y\s+Primitivo\s+de\s+Menfis[-\s]+Mizraim\s+y\s+de\s+la\s+Orden\s+de\s+los\s+Ritos\s+Unidos\s+de\s+Menfis\s*[&y]\s*Mizraim'
    new_name = 'Gran Logia Simbólica del Rito Antiguo y Primitivo de Menfis Misraim de la República del Paraguay'
    content = re.sub(pattern, new_name, content, flags=re.IGNORECASE)
    
    # Variante con "de la del"
    pattern2 = r'Gran\s+Logia\s+de\s+Espa[ñn]a\s+de\s+Menfis[-\s]+Mizraim\s+de\s+la\s+del\s+Rito\s+Antiguo\s+y\s+Primitivo\s+de\s+Menfis[-\s]+Mizraim\s+y\s+de\s+la\s+Orden\s+de\s+los\s+Ritos\s+Unidos\s+de\s+Menfis\s*[&y]\s*Mizraim'
    content = re.sub(pattern2, new_name, content, flags=re.IGNORECASE)
    
    # Variante más corta (solo "Gran Logia de España de Menfis-Mizraim" seguida de "del Rito")
    pattern3 = r'Gran\s+Logia\s+de\s+Espa[ñn]a\s+de\s+Menfis[-\s]+Mizraim\s+del\s+Rito\s+Antiguo\s+y\s+Primitivo'
    content = re.sub(pattern3, 'Gran Logia Simbólica del Rito Antiguo y Primitivo', content, flags=re.IGNORECASE)
    
    # Variante más corta aún
    pattern4 = r'Gran\s+Logia\s+de\s+Espa[ñn]a\s+de\s+Menfis[-\s]+Mizraim'
    content = re.sub(pattern4, 'Gran Logia Simbólica', content, flags=re.IGNORECASE)
    
    # Eliminar la referencia a "Ritos unificados por José GARIBALDI en 1881"
    pattern5 = r'Gran\s+Logia\s+Simbólica\s+del\s+Rito\s+Antiguo\s+y\s+Primitivo\s+de\s+Menfis\s+Misraim\s+de\s+la\s+República\s+del\s+Paraguay\s*[–\-—]\s*Ritos\s+(?:unificados|reunificados)\s+por\s+Jos[ée]\s+GARIBALDI\s+en\s+1881'
    content = re.sub(pattern5, 'Gran Logia Simbólica del Rito Antiguo y Primitivo de Menfis Misraim de la República del Paraguay', content, flags=re.IGNORECASE)
    
    # Reemplazar "en el Oriente de:………………..." por "en el Oriente del Paraguay"
    pattern6 = r'en\s+el\s+Oriente\s+de\s*[:\.]+\s*'
    content = re.sub(pattern6, 'en el Oriente del Paraguay ', content, flags=re.IGNORECASE)
    
    # Unir líneas fragmentadas innecesariamente
    # Detectar líneas que terminan sin puntuación final y unirlas con la siguiente
    lines = content.split('\n')
    joined_lines = []
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        # Si la línea no termina con puntuación final ni está vacía, unirla con la siguiente
        if line and not line.endswith(('.', '!', '?', ':', '…', '∴')) and not line.startswith('---') and i + 1 < len(lines):
            next_line = lines[i + 1].strip()
            # Unir si la siguiente línea no es un marcador de página ni está vacía
            if next_line and not next_line.startswith('--- PÁGINA') and not next_line.startswith('--- **PÁGINA'):
                joined_lines.append(line + ' ' + next_line)
                i += 2
                continue
        joined_lines.append(line)
        i += 1
    
    content = '\n'.join(joined_lines)
    
    # Limpiar secuencia redundante en portadas/encabezados
    # Patrón: Gran Logia Simbólica 2011 + RITO ANTIGUO... + ORDEN DE LOS RITOS... + Ritos reunificados... + Gran Logia Simbólica
    # Reemplazar por solo: RITO ANTIGUO Y PRIMITIVO DE MENFIS-MIZRAIM
    pattern7 = r'Gran\s+Logia\s+Simbólica\s+2011\s+RITO\s+ANTIGUO\s+Y\s+PRIMITIVO\s+DE\s+MENFIS-MIZRAIM\s+ORDEN\s+DE\s+LOS\s+RITOS\s+UNIDOS\s+DE\s+MENFIS\s*[&y]\s*MIZRAIM\s+Ritos\s+(?:reunificados|unificados)\s+por\s+Jos[ée]\s+GARIBALDI\s+en\s+1881\s+Gran\s+Logia\s+Simbólica'
    content = re.sub(pattern7, 'RITO ANTIGUO Y PRIMITIVO DE MENFIS-MIZRAIM', content, flags=re.IGNORECASE)
    
    # Eliminar "Gran Logia Simbólica 2011" (línea suelta)
    pattern8 = r'Gran\s+Logia\s+Simbólica\s+2011\s*\n?'
    content = re.sub(pattern8, '', content, flags=re.IGNORECASE)
    
    return content

def extract_pages(text):
    """Extract all pages with their content from the source text."""
    lines = text.split('\n')

    # First pass: identify recurring lines (headers/footers on 3+ pages)
    line_freq = {}
    for line in lines:
        stripped = line.strip()
        if stripped and not stripped.isdigit():
            line_freq[stripped] = line_freq.get(stripped, 0) + 1
    recurring = {l for l, c in line_freq.items() if c >= 3}

    skip_prefixes = [
        'GRAN LOGIA DE ESPANA DE MENFIS-MIZRAIM',
        'GRAN LOGIA DE ESPAÑA DE MENFIS-MIZRAIM',
        'A LA GLORIA DEL SUBLIME ARQUITECTO DE LOS MUNDOS',
        'A la Gloria del Sublime Arquitecto de los Mundos',
        'Bajo la Autoridad Espiritual',
        'Federacion de Logias',
        'Federación de Logias',
        'Traducidos del original',
        'RITUALES, CEREMONIAS E INSTRUCCIONES',
        'Rituales e Instrucciones de los',
        'Rituales, Ceremonias e Instrucciones',
        'HISTORIA. GRANDES CONSTITUCIONES',
        'HISTORIA.GRANDES CONSTITUCIONES',
    ]

    pages = {}
    current_page = None
    current_lines = []

    for line in lines:
        stripped = line.strip()

        if stripped.isdigit() and 2 <= int(stripped) <= 400:
            num = int(stripped)
            if current_page is not None and current_lines:
                pages[current_page] = current_lines
            current_page = num
            current_lines = []
            continue

        if current_page is not None:
            if not stripped:
                continue
            # Only skip if it's a known recurring header
            if stripped in recurring:
                skip_it = False
                for pref in skip_prefixes:
                    if stripped.startswith(pref):
                        skip_it = True
                        break
                if skip_it:
                    continue
            # Skip purely decorative lines
            if re.match(r'^[\s\u00b7\u2022\u2023\u25cf\u25d8\u2713\u2714\u2666\u2663\u2605\u2606\u2234\*\-\.\~\,\s]+$', stripped):
                continue
            current_lines.append(stripped)

    if current_page is not None and current_lines:
        pages[current_page] = current_lines

    return pages

SKIP_BOLD = {'Y','E','O','A','DE','DEL','EN','LA','EL','LOS','LAS','UN','UNA',
             'POR','CON','QUE','ES','NO','SE','SU','AL','LO','LE','SUS','HA',
             'HE','HOY','VA','DA','SAN','ERA','SON','HAN','SER','MÁS',
             'I','III','II','IV','VI','VII','VIII','IX','XI','XII','XIII',
             '1ER','2º','3ER','2O','D','L','M','H','V','VV','MM','HH',
             'AAP','CC','TT','S','A','N','O','P','R','B','C','J','K',
             'G','T','Hº','H','S','N','O','R','F','P','J','B','Z'}

def has_uppercase_word(word):
    """Check if word qualifies for bolding (3+ alpha chars, all uppercase/accents, not in skip list)."""
    word_clean = word.strip('.,;:!?¿¡"\'()[]{}«»-–—')
    if not word_clean or len(word_clean) < 3:
        return False
    if word_clean in SKIP_BOLD:
        return False
    # Must have at least one alpha character (to avoid bolding numbers like 1881)
    alpha_chars = [c for c in word_clean if c.isalpha()]
    if not alpha_chars:
        return False
    return all(c.isupper() for c in alpha_chars)

def mark_bold_html(text):
    words = text.split()
    result = []
    for w in words:
        if has_uppercase_word(w):
            prefix = ''
            suffix = ''
            clean = w
            while clean and not clean[-1].isalpha():
                suffix = clean[-1] + suffix
                clean = clean[:-1]
            while clean and not clean[0].isalpha():
                prefix = prefix + clean[0]
                clean = clean[1:]
            result.append(f'{prefix}<strong>{clean}</strong>{suffix}')
        else:
            result.append(w)
    return ' '.join(result)

def mark_bold_md(text):
    words = text.split()
    result = []
    for w in words:
        if has_uppercase_word(w):
            prefix = ''
            suffix = ''
            clean = w
            while clean and not clean[-1].isalpha():
                suffix = clean[-1] + suffix
                clean = clean[:-1]
            while clean and not clean[0].isalpha():
                prefix = prefix + clean[0]
                clean = clean[1:]
            result.append(f'{prefix}**{clean}**{suffix}')
        else:
            result.append(w)
    return ' '.join(result)

# Speaker patterns for ritual dialog (V∴M∴, 1er Vig∴, etc.)
SPEAKER_SET = {
    'V∴M∴','V∴M','V∴ M∴','V∴ M','M∴ de C∴','M∴ de  C∴','M∴D∴C∴','M∴D∴C∴I∴',
    'V∴ M∴de C∴','V∴ M∴ de C∴',
    '1er Vig∴','1er  Vig∴','E∴V∴','2º Vig∴','2º  Vig∴','2º Vig∴ Electo',
    '1er Vig∴ Electo','V∴1er Vig∴','V∴2º Vig∴',
    'Orad∴','Or∴','V∴Orador','V∴Orador :',
    'Sec∴','V∴Secretario','Tes∴','Hos∴','Exp∴','Experto∴',
    'Cand∴','G∴T∴','V∴G∴T∴','V∴Hierofante',
    'Col∴de Arm∴','Col∴ de Ar∴','Col∴ de  Ar∴',
    '1er Inst∴','2º Inst∴','Expert Inst∴','Experto Inst∴',
    'Los 2 VVig∴','TT∴ los OOf∴','MM∴MM∴',
    'V∴M∴ Electo','V∴M ∴','V∴M∴ :','M∴ de C∴:',
    'Todos los HH∴',
}

SPEAKER_RE = re.compile(r'^(' + '|'.join(re.escape(s) for s in sorted(SPEAKER_SET, key=len, reverse=True)) + r')\s*[:\u2794]?\s*$')

def is_speaker_line(s):
    s = s.strip()
    if not s:
        return False
    return bool(SPEAKER_RE.match(s))

def group_dialog_lines(content_lines):
    """Group lines into segments: regular lines and dialog pairs (speaker, speech_lines)."""
    segments = []
    i = 0
    while i < len(content_lines):
        s = content_lines[i].strip()
        if is_speaker_line(s):
            speaker = s
            speech_lines = []
            i += 1
            while i < len(content_lines):
                ns = content_lines[i].strip()
                if not ns:
                    i += 1
                    continue
                if is_speaker_line(ns):
                    break  # next speaker starts a new pair
                speech_lines.append(ns)
                i += 1
            segments.append(('dialog', speaker, speech_lines))
        else:
            segments.append(('line', s))
            i += 1
    return segments


def format_line_html(line):
    """Classify and format a single line for HTML."""
    s = line.strip()
    if not s:
        return '', 'empty'

    if s in ['CATECISMO'] or s.startswith('INSTRUCCION DEL') or s.startswith('INSTRUCCIÓN DEL'):
        return f'<div class="catecismo">{s}</div>', 'catecismo'
    if s == 'ÍNDICE' or s == 'INDICE':
        return f'<div class="indice">{s}</div>', 'indice'
    if s.startswith('Pregunta') or s.startswith('¿Pregunta'):
        q = s.replace('Pregunta', '<strong>Pregunta</strong>')
        q = mark_bold_html(q)
        return f'<div class="question">{q}</div>', 'question'
    if s.startswith('Respuesta'):
        a = s.replace('Respuesta', '<strong>Respuesta</strong>')
        a = mark_bold_html(a)
        return f'<div class="answer">{a}</div>', 'answer'
    if 'A LA GLORIA' in s or 'A la Gloria' in s:
        return f'<div class="gloria">{mark_bold_html(s)}</div>', 'gloria'
    if 'RITO ANTIGUO Y PRIMITIVO' in s:
        return f'<div class="rito">{mark_bold_html(s)}</div>', 'rito'
    if 'ORDEN DE LOS RITOS UNIDOS' in s:
        return f'<div class="orden">{mark_bold_html(s)}</div>', 'orden'
    if s.startswith('GRAN LOGIA'):
        return f'<div class="granlogia">{mark_bold_html(s)}</div>', 'granlogia'
    if 'Ritos reunificados' in s or 'Ritos unificados' in s:
        return f'<div class="reunificados">{mark_bold_html(s)}</div>', 'reunificados'
    if s.startswith('PROCLAMACIÓN') or s.startswith('PROCLAMACION'):
        return f'<div class="proclamation">{mark_bold_html(s)}</div>', 'proclamation'
    if s.startswith('CAPÍTULO') or s.startswith('CAPITULO') or s.startswith('CAP.'):
        return f'<div class="chapter">{mark_bold_html(s)}</div>', 'chapter'
    if s.startswith('TÍTULO') or s.startswith('TITULO'):
        return f'<div class="title-header">{mark_bold_html(s)}</div>', 'title'
    if s.startswith('Art.') or s.startswith('Artículo') or s.startswith('ARTICULO') or s.startswith('ART.'):
        return f'<div class="article">{mark_bold_html(s)}</div>', 'article'
    if s.startswith('DISPOSICIONES'):
        return f'<div class="chapter">{mark_bold_html(s)}</div>', 'chapter'
    if re.match(r'^[\(\)\u2794\*].*[\)]$', s):
        return f'<div class="stage">{s}</div>', 'stage'
    if s.startswith('CHEQUETET') or s.startswith('MAK-BENAH') or s.startswith('ARELICH') or s.startswith('VOMALITES'):
        return f'<div class="aclamacion">{s}</div>', 'aclamacion'
    if s.startswith('Amén') or s.startswith('Soberano Maestro') or s.startswith('Eterno Regenerador') or s.startswith('¡Soberano'):
        return f'<div class="oracion">{mark_bold_html(s)}</div>', 'oracion'
    if re.match(r'^\d+\.\s+', s) and len(s) < 80:
        return f'<div class="article">{mark_bold_html(s)}</div>', 'article'

    return f'<div class="text-block">{mark_bold_html(s)}</div>', 'text'


def generate_html(pages_dict, doc_config):
    page_nums = doc_config['pages']
    title = doc_config['title']

    lines = []
    lines.append('<!DOCTYPE html>')
    lines.append('<html lang="es">')
    lines.append('<head>')
    lines.append('<meta charset="UTF-8">')
    lines.append('<meta name="viewport" content="width=device-width, initial-scale=1.0">')
    lines.append(f'<title>{title}</title>')
    lines.append('''<style>
* { margin: 0; padding: 0; box-sizing: border-box; }

body {
  font-family: 'Times New Roman', Times, serif;
  background: #3a3226;
  color: #1a1a1a;
  padding: 0;
}

.container {
  max-width: 800px;
  margin: 0 auto;
  background: #faf8f2;
  min-height: 100vh;
}

.page {
  padding: 80px 100px 60px;
  page-break-after: always;
  border-bottom: 1px solid #d4c9b0;
}

.page-marker {
  text-align: center;
  margin: 0 0 30px;
  padding: 4px 0;
  font-size: 10px;
  color: #c4b69a;
  letter-spacing: 6px;
}
.page-marker::before { content: "~ ~ ~"; }

.gloria {
  font-size: 13px;
  font-variant: small-caps;
  letter-spacing: 3px;
  color: #8b7355;
  margin: 30px 0;
  text-align: center;
  line-height: 1.8;
}

.rito {
  font-size: 15px;
  font-weight: bold;
  text-transform: uppercase;
  letter-spacing: 4px;
  margin: 20px 0;
  text-align: center;
  color: #1a1a1a;
  line-height: 1.6;
}

.orden {
  font-size: 12px;
  font-style: italic;
  color: #555;
  margin: 20px 0;
  text-align: center;
  line-height: 1.6;
}

.granlogia {
  font-size: 12px;
  font-weight: bold;
  color: #8b7355;
  margin: 15px 0;
  text-align: center;
  letter-spacing: 2px;
}

.reunificados {
  font-size: 12px;
  font-style: italic;
  text-align: center;
  margin: 15px 0;
  color: #666;
  line-height: 1.6;
}

.text-block {
  margin: 6px 0;
  line-height: 1.7;
  text-align: justify;
  font-size: 12pt;
}

/* Headers */
.proclamation {
  font-size: 13px;
  font-weight: bold;
  text-align: center;
  margin: 25px 0 20px;
  letter-spacing: 2px;
  text-transform: uppercase;
}

.chapter {
  font-size: 14px;
  font-weight: bold;
  margin: 30px 0 15px;
  text-align: center;
  letter-spacing: 1px;
}

.title-header {
  font-size: 13px;
  font-weight: bold;
  font-variant: small-caps;
  margin: 20px 0 10px;
  letter-spacing: 1px;
}

.article {
  margin: 5px 0 5px 15px;
  line-height: 1.6;
}

.catecismo {
  font-size: 16px;
  font-weight: bold;
  text-align: center;
  margin: 30px 0 20px;
  font-variant: small-caps;
  letter-spacing: 3px;
}

.indice {
  font-size: 18px;
  font-weight: bold;
  text-align: center;
  margin: 40px 0 25px;
  letter-spacing: 5px;
}

.question {
  margin: 18px 0 4px;
  font-style: italic;
  color: #2c2c2c;
  line-height: 1.6;
}

.answer {
  margin: 0 0 18px 25px;
  padding-left: 12px;
  border-left: 2px solid #c4b69a;
  color: #333;
  line-height: 1.6;
}

.stage {
  font-style: italic;
  color: #666;
  margin: 12px 0;
  padding-left: 25px;
  line-height: 1.6;
}

.aclamacion {
  text-align: center;
  font-weight: bold;
  margin: 20px 0;
  letter-spacing: 2px;
  font-size: 13px;
}

.oracion {
  text-align: center;
  font-style: italic;
  margin: 15px 0;
  color: #444;
  line-height: 1.6;
}

/* Decorative separator for section pages */
.section-sep {
  text-align: center;
  margin: 30px 0;
  color: #8b7355;
  font-size: 18px;
  letter-spacing: 10px;
}

@media print {
  body { background: white; }
  .container { max-width: none; }
  .page { padding: 60px 80px; border: none; }
}

.cover-image { width: 100%; max-width: 100%; display: block; margin: 0 auto; }
.toc-image { width: 100%; max-width: 100%; display: block; margin: 0 auto; }
.diagram-image { max-width: 90%; display: block; margin: 15px auto; }

.dialog-table { width: 100%; border: none; border-collapse: collapse; margin: 4px 0; }
.dialog-table td { vertical-align: top; padding: 2px 6px; }
.speaker { width: 22%; font-weight: bold; white-space: nowrap; text-align: left; padding-left: 0 !important; }
.speech { width: 78%; text-align: justify; line-height: 1.6; }
.speech br { content: ''; display: block; margin: 4px 0; }

/* Navigable Index Styles */
.nav-index {
  padding: 60px 80px;
  background: #faf8f2;
}

.indice-title {
  font-size: 20px;
  font-weight: bold;
  text-align: center;
  margin-bottom: 40px;
  color: #1a1a1a;
  letter-spacing: 4px;
  font-variant: small-caps;
}

.index-section {
  margin-bottom: 30px;
}

.index-section h3 {
  font-size: 14px;
  font-weight: bold;
  color: #8b7355;
  margin-bottom: 12px;
  border-bottom: 1px solid #c4b69a;
  padding-bottom: 6px;
  text-transform: uppercase;
  letter-spacing: 2px;
}

.index-section ul {
  list-style: none;
  padding-left: 0;
  margin: 0;
}

.index-section li {
  margin: 8px 0;
  line-height: 1.5;
}

.index-section a {
  color: #2c2c2c;
  text-decoration: none;
  font-size: 11pt;
  transition: color 0.2s;
}

.index-section a:hover {
  color: #8b7355;
  text-decoration: underline;
}

.index-section a::before {
  content: "▸ ";
  color: #c4b69a;
}
</style>''')
    lines.append('</head>')
    lines.append('<body>')
    lines.append('<div class="container">')

    is_completo = doc_config['id'] == 'Rituales_Simbolicos_Completo'
    img_dir = os.path.join(os.path.dirname(__file__), 'images')

    # Insert cover image for complete work
    if is_completo:
        cover_path = os.path.join(img_dir, 'cover.png')
        if os.path.exists(cover_path):
            # Use data URI for portability
            import base64
            with open(cover_path, 'rb') as f:
                b64 = base64.b64encode(f.read()).decode()
            lines.append(f'<img class="cover-image" src="data:image/png;base64,{b64}" alt="Portada">')
            lines.append('<div style="page-break-after:always;"></div>')
        
        # TOC image
        toc_path = os.path.join(img_dir, 'page4_img_2.png')
        if os.path.exists(toc_path):
            with open(toc_path, 'rb') as f:
                b64 = base64.b64encode(f.read()).decode()
            lines.append(f'<img class="toc-image" src="data:image/png;base64,{b64}" alt="Índice">')
            lines.append('<div style="page-break-after:always;"></div>')
        
        # Add navigable index
        lines.append('<div class="nav-index">')
        lines.append('<h2 class="indice-title">ÍNDICE NAVEGABLE</h2>')
        lines.append('<div class="index-section">')
        lines.append('<h3>Primer Grado - Aprendiz</h3>')
        lines.append('<ul>')
        lines.append('<li><a href="#page4">Esquema de una Logia en 1er grado</a></li>')
        lines.append('<li><a href="#page5">Ritual Largo del 1er grado Aprendiz - Introducción</a></li>')
        lines.append('<li><a href="#page12">Apertura de los trabajos en 1er grado Aprendiz</a></li>')
        lines.append('<li><a href="#page24">Cierre de los trabajos en 1er grado Aprendiz</a></li>')
        lines.append('<li><a href="#page31">Ritual Corto del 1er grado Aprendiz</a></li>')
        lines.append('<li><a href="#page46">Ritual de Iniciación al 1er grado Aprendiz - Cámara de Reflexión</a></li>')
        lines.append('<li><a href="#page47">Ritual Ceremonia de Iniciación al 1er grado Aprendiz</a></li>')
        lines.append('<li><a href="#page68">Instrucción del 1er grado Aprendiz - Catecismo</a></li>')
        lines.append('<li><a href="#page76">Ritual de los Trabajos de un Triángulo Masónico en 1er grado</a></li>')
        lines.append('<li><a href="#page82">Ritual de Audiencia a un Profano - Pase bajo la Venda</a></li>')
        lines.append('<li><a href="#page87">Juramento de Afiliación</a></li>')
        lines.append('<li><a href="#page88">Ritual de Afiliación para HH∴ de nuestro Rito</a></li>')
        lines.append('<li><a href="#page94">Ritual de Afiliación para HH∴ de distinto Rito</a></li>')
        lines.append('</ul>')
        lines.append('</div>')
        lines.append('<div class="index-section">')
        lines.append('<h3>Segundo Grado - Compañero</h3>')
        lines.append('<ul>')
        lines.append('<li><a href="#page101">Ritual completo 2º grado Compañero</a></li>')
        lines.append('<li><a href="#page102">Apertura de los trabajos en 2º grado Compañero</a></li>')
        lines.append('<li><a href="#page108">Cierre de los trabajos en 2º grado Compañero</a></li>')
        lines.append('<li><a href="#page112">Preliminares Ceremonia de Recepción al 2º grado Compañero</a></li>')
        lines.append('<li><a href="#page113">Ritual Ceremonia de Recepción al 2º grado Compañero</a></li>')
        lines.append('<li><a href="#page133">Instrucción del 2º grado Compañero - Catecismo</a></li>')
        lines.append('</ul>')
        lines.append('</div>')
        lines.append('<div class="index-section">')
        lines.append('<h3>Tercer Grado - Maestro</h3>')
        lines.append('<ul>')
        lines.append('<li><a href="#page137">Ritual completo 3er grado Maestro</a></li>')
        lines.append('<li><a href="#page139">Apertura de los trabajos en 3er grado Maestro</a></li>')
        lines.append('<li><a href="#page147">Cierre de los trabajos en 3er grado Maestro</a></li>')
        lines.append('<li><a href="#page151">Preliminares Ceremonia de Recepción al 3er grado Maestro</a></li>')
        lines.append('<li><a href="#page153">Ritual Ceremonia de Recepción al 3er grado Maestro</a></li>')
        lines.append('<li><a href="#page169">Instrucción del 3er grado Maestro - Catecismo</a></li>')
        lines.append('</ul>')
        lines.append('</div>')
        lines.append('<div class="index-section">')
        lines.append('<h3>Otros Rituales</h3>')
        lines.append('<ul>')
        lines.append('<li><a href="#page178">Ritual de Fundación de un Triángulo Masónico</a></li>')
        lines.append('<li><a href="#page189">Ritual de Fundación de una Logia Masónica</a></li>')
        lines.append('<li><a href="#page198">Ritual de Elecciones e Instalación del V∴ M∴ y de los Oficiales</a></li>')
        lines.append('<li><a href="#page209">Ritual de Tenida Blanca Cerrada</a></li>')
        lines.append('<li><a href="#page212">Ritual de Tenida Blanca Abierta</a></li>')
        lines.append('<li><a href="#page217">Ritual de Reconocimiento Conyugal Masónico</a></li>')
        lines.append('<li><a href="#page231">Ritual de Adopción de un Lobatón</a></li>')
        lines.append('<li><a href="#page241">Ritual de San Juan de Verano dentro del Templo</a></li>')
        lines.append('<li><a href="#page248">Ritual de San Juan de Verano en el exterior</a></li>')
        lines.append('<li><a href="#page256">Ritual de Mesa</a></li>')
        lines.append('<li><a href="#page264">Rituales Fúnebres Masónicos</a></li>')
        lines.append('<li><a href="#page282">Historia de la Orden de los Ritos Unidos de Menfis & Mizraim</a></li>')
        lines.append('<li><a href="#page301">Grandes Hierofantes</a></li>')
        lines.append('<li><a href="#page321">Grandes Constituciones y Reglamentos Generales</a></li>')
        lines.append('</ul>')
        lines.append('</div>')
        lines.append('</div>')
        lines.append('<div style="page-break-after:always;"></div>')

    start_page = 5 if is_completo else 2

    for idx, pg in enumerate(page_nums):
        if pg < start_page:
            continue

        if pg not in pages_dict:
            continue

        # Add ID for navigation
        lines.append(f'<div class="page" id="page{pg}">')

        # Group into dialog pairs
        segments = group_dialog_lines(pages_dict[pg])
        for seg in segments:
            if seg[0] == 'dialog':
                speaker = mark_bold_html(seg[1])
                speech_text = '<br>'.join(mark_bold_html(l) for l in seg[2])
                lines.append(f'<table class="dialog-table"><tr><td class="speaker">{speaker}</td><td class="speech">{speech_text}</td></tr></table>')
            else:
                html_line, _ = format_line_html(seg[1])
                if html_line:
                    lines.append(html_line)

        lines.append('</div>')

    lines.append('</div>')
    lines.append('</body>')
    lines.append('</html>')
    return '\n'.join(lines)


def format_line_md(line):
    s = line.strip()
    if not s:
        return '', 'empty'
    if s in ['CATECISMO'] or s.startswith('INSTRUCCION DEL') or s.startswith('INSTRUCCIÓN DEL') or s == 'ÍNDICE' or s == 'INDICE':
        return f'## {s}', 'header'
    if s.startswith('Pregunta') or s.startswith('¿Pregunta'):
        return mark_bold_md(s.replace('Pregunta', '**Pregunta:**')), 'question'
    if s.startswith('Respuesta'):
        return f'> {mark_bold_md(s.replace("Respuesta", "**Respuesta:**"))}', 'answer'
    if 'A LA GLORIA' in s or 'A la Gloria' in s:
        return f'*{mark_bold_md(s)}*', 'italic'
    if 'RITO ANTIGUO' in s:
        return f'# {mark_bold_md(s)}', 'big'
    if 'ORDEN DE LOS RITOS UNIDOS' in s:
        return f'*{mark_bold_md(s)}*', 'italic'
    if s.startswith('GRAN LOGIA'):
        return f'**{mark_bold_md(s)}**', 'bold'
    if 'Ritos reunificados' in s or 'Ritos unificados' in s:
        return f'*{mark_bold_md(s)}*', 'italic'
    if s.startswith('PROCLAMACIÓN') or s.startswith('PROCLAMACION') or s.startswith('CAPÍTULO') or s.startswith('CAPITULO') or s.startswith('DISPOSICIONES'):
        return f'## {mark_bold_md(s)}', 'header'
    if s.startswith('TÍTULO') or s.startswith('TITULO'):
        return f'### {mark_bold_md(s)}', 'header'
    if s.startswith('CHEQUETET') or s.startswith('MAK-BENAH') or s.startswith('ARELICH') or s.startswith('VOMALITES'):
        return f'**{s}**', 'bold'
    if s.startswith('Amén') or s.startswith('Soberano Maestro') or s.startswith('Eterno Regenerador'):
        return f'*{mark_bold_md(s)}*', 'italic'
    if re.match(r'^[\(\)\u2794\*].*[\)]$', s):
        return f'*{s}*', 'italic'
    return mark_bold_md(s), 'text'


def generate_md(pages_dict, doc_config):
    page_nums = doc_config['pages']
    title = doc_config['title']
    lines = [f'# {title}', '']
    start_page = 5 if doc_config['id'] == 'Rituales_Simbolicos_Completo' else 2

    for idx, pg in enumerate(page_nums):
        if pg < start_page:
            continue
        lines.append('')
        lines.append(f'--- **PÁGINA {pg}** ---')
        lines.append('')
        if pg not in pages_dict:
            lines.append(f'*[página {pg} - contenido no disponible en esta sección]*')
            lines.append('')
            continue
        segments = group_dialog_lines(pages_dict[pg])
        for seg in segments:
            if seg[0] == 'dialog':
                speaker = mark_bold_md(seg[1])
                speech_text = '  \n'.join(mark_bold_md(l) for l in seg[2])
                lines.append(f'| {speaker} | {speech_text} |')
            else:
                md_line, _ = format_line_md(seg[1])
                if md_line:
                    lines.append(md_line)
                lines.append('')
    return '\n'.join(lines)


def add_para(doc, text, bold=False, italic=False, alignment=None, font_size=None, indent=None, auto_bold=True, heading_level=0):
    if heading_level:
        p = doc.add_heading(text, level=heading_level)
        if alignment:
            p.alignment = alignment
        return p

    p = doc.add_paragraph()
    if alignment:
        p.alignment = alignment
    if indent:
        p.paragraph_format.left_indent = indent

    if auto_bold and not bold:
        words = text.split()
        for i, w in enumerate(words):
            if has_uppercase_word(w):
                prefix = ''
                suffix = ''
                clean = w
                while clean and not clean[-1].isalpha():
                    suffix = clean[-1] + suffix
                    clean = clean[:-1]
                while clean and not clean[0].isalpha():
                    prefix = prefix + clean[0]
                    clean = clean[1:]
                if prefix:
                    r = p.add_run(prefix)
                    r.font.name = 'Times New Roman'
                    if italic: r.italic = True
                r = p.add_run(clean)
                r.bold = True
                r.font.name = 'Times New Roman'
                if italic: r.italic = True
                if font_size: r.font.size = font_size
                if suffix:
                    r = p.add_run(suffix)
                    r.font.name = 'Times New Roman'
                    if italic: r.italic = True
            else:
                r = p.add_run(w)
                r.font.name = 'Times New Roman'
                if italic: r.italic = True
                if font_size: r.font.size = font_size
            if i < len(words) - 1:
                r = p.add_run(' ')
                r.font.name = 'Times New Roman'
    else:
        run = p.add_run(text)
        run.bold = bold
        run.italic = italic
        run.font.name = 'Times New Roman'
        if font_size:
            run.font.size = font_size
    return p


def generate_docx(pages_dict, doc_config):
    page_nums = doc_config['pages']
    title = doc_config['title']
    is_completo = doc_config['id'] == 'Rituales_Simbolicos_Completo'

    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(11)
    style.paragraph_format.line_spacing = 1.5

    # Configure heading styles
    for level in range(1, 5):
        hs = doc.styles[f'Heading {level}']
        hs.font.name = 'Times New Roman'
        hs.font.color.rgb = RGBColor(0x1a, 0x1a, 0x1a)

    # Add automatic page numbering in footer
    section = doc.sections[0]
    section.different_first_page_header_footer = True
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fp.add_run()
    fldChar1 = run._r.makeelement('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fldChar', {})
    fldChar1.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fldCharType', 'begin')
    run._r.append(fldChar1)

    run2 = fp.add_run()
    instrText = run2._r.makeelement('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}instrText', {})
    instrText.text = ' PAGE '
    run2._r.append(instrText)

    run3 = fp.add_run()
    fldChar2 = run3._r.makeelement('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fldChar', {})
    fldChar2.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fldCharType', 'end')
    run3._r.append(fldChar2)

    # Title page (no header/footer on first page)
    add_para(doc, title, bold=True, font_size=Pt(18), alignment=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()  # blank line
    doc.add_page_break()

    start_page = 5 if is_completo else 2

    for idx, pg in enumerate(page_nums):
        if pg < start_page:
            continue

        if pg not in pages_dict:
            add_para(doc, f'[página {pg} - contenido no disponible]', italic=True)
            doc.add_page_break()
            continue

        segments = group_dialog_lines(pages_dict[pg])
        for seg in segments:
            if seg[0] == 'dialog':
                speaker_text = seg[1]
                speech_list = seg[2]
                # Two-column table without borders
                table = doc.add_table(rows=1, cols=2)
                # Remove borders
                for row in table.rows:
                    for cell in row.cells:
                        tc = cell._tc
                        tcPr = tc.get_or_add_tcPr()
                        tcBorders = OxmlElement('w:tcBorders')
                        for bn in ['top', 'left', 'bottom', 'right']:
                            b = OxmlElement(f'w:{bn}')
                            b.set(qn('w:val'), 'nil')
                            tcBorders.append(b)
                        tcPr.append(tcBorders)
                # Left cell: speaker
                cell_l = table.cell(0, 0)
                cell_l.width = Inches(1.5)
                cp = cell_l.paragraphs[0]
                run = cp.add_run(speaker_text)
                run.bold = True
                run.font.name = 'Times New Roman'
                run.font.size = Pt(10)
                # Right cell: speech lines
                cell_r = table.cell(0, 1)
                for si, sl in enumerate(speech_list):
                    if si > 0:
                        cell_r.add_paragraph()
                    cp = cell_r.paragraphs[-1]
                    words = sl.split()
                    for wi, w in enumerate(words):
                        if has_uppercase_word(w):
                            prefix = ''
                            suffix = ''
                            clean = w
                            while clean and not clean[-1].isalpha():
                                suffix = clean[-1] + suffix
                                clean = clean[:-1]
                            while clean and not clean[0].isalpha():
                                prefix = prefix + clean[0]
                                clean = clean[1:]
                            if prefix:
                                r = cp.add_run(prefix)
                                r.font.name = 'Times New Roman'
                                r.font.size = Pt(10)
                            r = cp.add_run(clean)
                            r.bold = True
                            r.font.name = 'Times New Roman'
                            r.font.size = Pt(10)
                            if suffix:
                                r = cp.add_run(suffix)
                                r.font.name = 'Times New Roman'
                                r.font.size = Pt(10)
                        else:
                            r = cp.add_run(w)
                            r.font.name = 'Times New Roman'
                            r.font.size = Pt(10)
                        if wi < len(words) - 1:
                            r = cp.add_run(' ')
                            r.font.name = 'Times New Roman'
            else:
                s = seg[1]

                # Heading level detection for TOC
                h = 0
                if s in ['CATECISMO'] or s.startswith('INSTRUCCION DEL') or s.startswith('INSTRUCCIÓN DEL'):
                    h = 2
                elif s == 'ÍNDICE' or s == 'INDICE':
                    h = 1
                elif s.startswith('PROCLAMACIÓN') or s.startswith('PROCLAMACION'):
                    h = 1
                elif s.startswith('CAPÍTULO') or s.startswith('CAPITULO'):
                    h = 2
                elif s.startswith('TÍTULO') or s.startswith('TITULO'):
                    h = 3
                elif s.startswith('DISPOSICIONES'):
                    h = 2
                elif s.startswith('GRAN LOGIA') and '2011' in s:
                    h = 1
                elif 'RITO ANTIGUO Y PRIMITIVO' in s:
                    h = 1
                elif 'ORDEN DE LOS RITOS UNIDOS' in s and is_completo:
                    h = 2

                if h:
                    add_para(doc, s, bold=True, heading_level=h)
                elif s.startswith('Pregunta') or s.startswith('¿Pregunta'):
                    add_para(doc, s.replace('Pregunta', 'Pregunta:'), italic=True)
                elif s.startswith('Respuesta'):
                    add_para(doc, s.replace('Respuesta', 'Respuesta:'), indent=Inches(0.4))
                elif 'A LA GLORIA' in s or 'A la Gloria' in s:
                    add_para(doc, s, italic=True, font_size=Pt(11), alignment=WD_ALIGN_PARAGRAPH.CENTER)
                elif 'RITO ANTIGUO' in s:
                    add_para(doc, s, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, auto_bold=False)
                elif 'ORDEN DE LOS RITOS UNIDOS' in s:
                    add_para(doc, s, italic=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
                elif s.startswith('GRAN LOGIA'):
                    add_para(doc, s, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, auto_bold=False)
                elif 'Ritos reunificados' in s or 'Ritos unificados' in s:
                    add_para(doc, s, italic=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
                elif s.startswith('PROCLAMACIÓN') or s.startswith('PROCLAMACION') or s.startswith('CAPÍTULO') or s.startswith('CAPITULO'):
                    add_para(doc, s, bold=True, font_size=Pt(13), alignment=WD_ALIGN_PARAGRAPH.CENTER, auto_bold=False)
                elif s.startswith('TÍTULO') or s.startswith('TITULO') or s.startswith('DISPOSICIONES'):
                    add_para(doc, s, bold=True, auto_bold=False)
                elif s.startswith('CHEQUETET') or s.startswith('MAK-BENAH') or s.startswith('ARELICH'):
                    add_para(doc, s, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, auto_bold=False)
                elif s.startswith('Amén') or s.startswith('Soberano Maestro'):
                    add_para(doc, s, italic=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
                elif re.match(r'^[\(\)\u2794\*].*[\)]$', s):
                    add_para(doc, s, italic=True, auto_bold=False)
                elif s.startswith('Art.') or s.startswith('Artículo') or s.startswith('ART.'):
                    add_para(doc, s, indent=Inches(0.2))
                else:
                    add_para(doc, s)

        # Page break between document pages
        if idx < len(page_nums) - 1:
            doc.add_page_break()

    return doc


def main():
    print("Cargando fuente completa...")
    source_text = load_source()

    print("Extrayendo páginas...")
    all_pages = extract_pages(source_text)
    print(f"  {len(all_pages)} páginas extraídas: {min(all_pages.keys()) if all_pages else '?'} - {max(all_pages.keys()) if all_pages else '?'}")

    html_dir = os.path.join(OUTPUT, 'html')
    md_dir = os.path.join(OUTPUT, 'md')
    doc_dir = os.path.join(OUTPUT, 'doc')
    for d in [html_dir, md_dir, doc_dir]:
        os.makedirs(d, exist_ok=True)

    for doc_cfg in DOCUMENTS:
        print(f'\n{doc_cfg["title"]}')

        # HTML
        html = generate_html(all_pages, doc_cfg)
        hpath = os.path.join(html_dir, f'{doc_cfg["id"]}.html')
        with open(hpath, 'w', encoding='utf-8') as f:
            f.write(html)
        print(f'  HTML: {doc_cfg["id"]}.html')

        # MD
        md = generate_md(all_pages, doc_cfg)
        mpath = os.path.join(md_dir, f'{doc_cfg["id"]}.md')
        with open(mpath, 'w', encoding='utf-8') as f:
            f.write(md)
        print(f'  MD:   {doc_cfg["id"]}.md')

        # DOCX
        doc = generate_docx(all_pages, doc_cfg)
        docpath = os.path.join(doc_dir, f'{doc_cfg["id"]}.docx')
        doc.save(docpath)
        print(f'  DOCX: {doc_cfg["id"]}.docx')

    print('\n¡Todos generados desde la misma fuente!')


if __name__ == '__main__':
    main()
