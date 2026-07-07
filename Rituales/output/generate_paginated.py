import os, re, sys
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUTPUT = r'C:\Temp\kitian\GITHUT\RItoMemphisMisraim\Rituales\output'

# ---- CONFIG: each source file, its pages, and how to detect page breaks ----
SOURCES = [
    {
        'id': 'Ritual_1_Instruccion_1er_Grado_Aprendiz',
        'title': 'Instrucción del 1er Grado - Aprendiz',
        'file': os.path.join(os.environ['TEMP'], 'pdf_imgs', 'ritual_1_completo.txt'),
        'pages': [68, 70, 71, 73, 74, 75],
        'page_header_lines': 2,  # lines after page number that are skip headers
    },
    {
        'id': 'Ritual_2_Instruccion_2o_Grado_Companero',
        'title': 'Instrucción del 2º Grado - Compañero',
        'file': os.path.join(os.environ['TEMP'], 'pdf_imgs_2', 'ritual_2_completo.txt'),
        'pages': [133, 134, 135, 136],
        'page_header_lines': 2,
    },
    {
        'id': 'Ritual_3_Maestro',
        'title': 'Ritual e Instrucción del 3er Grado - Maestro',
        'file': os.path.join(os.environ['TEMP'], 'pdf_imgs_2', 'ritual_3_completo.txt'),
        'pages': [137, 138, 147, 148, 149, 150, 151, 152, 153, 154, 155, 156,
                  139, 157, 158, 159, 160, 161, 162, 163, 164, 165, 166,
                  140, 167, 168, 169, 170, 171, 172, 173, 174, 175, 176,
                  141, 177, 142, 143, 144, 145, 146],
        'page_header_lines': 2,
    },
    {
        'id': 'Grandes_Constituciones',
        'title': 'Grandes Constituciones y Reglamentos Generales',
        'file': os.path.join(os.environ['TEMP'], 'pdf_text_new', 'ritual_0_completo.txt'),
        'pages': list(range(321, 391)),
        'page_header_lines': 3,
    },
    {
        'id': 'Rituales_Simbolicos_Completo',
        'title': 'Rituales e Instrucciones de los Grados Simbólicos - Obra Completa',
        'file': os.path.join(os.environ['TEMP'], 'pdf_text_new', 'ritual_4_completo.txt'),
        'pages': list(range(2, 391)),
        'page_header_lines': 4,
    },
]

SKIP_LINES = [
    'RITUALES, CEREMONIAS E INSTRUCCIONES',
    'HISTORIA. GRANDES CONSTITUCIONES',
    'HISTORIA.GRANDES CONSTITUCIONES',
    'GRAN LOGIA DE ESPANA DE MENFIS-MIZRAIM',
    'GRAN LOGIA DE ESPAÑA DE MENFIS-MIZRAIM',
    'LONSEO', 'ONSEO', 'CONSEIO', 'CONSEJO', 'CONSEJO',
    'S O', '. G', 'S CO S', 'UC O', 'S', 'G', 'OS G', 'S.',
]

def is_skip_header(line):
    s = line.strip()
    if not s:
        return True
    for pat in SKIP_LINES:
        if s.startswith(pat):
            return True
    return False

def load_text(path):
    with open(path, 'r', encoding='utf-8') as f:
        return f.read()

def split_into_pages(text, config):
    """Split text into pages based on page number markers."""
    lines = text.split('\n')
    pages = []
    current_page_num = None
    current_lines = []
    
    # Build list of page numbers to detect
    page_nums = set(str(p) for p in config['pages'])
    
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        
        # Check if this line is a page number
        if stripped in page_nums and (i == 0 or not lines[i-1].strip() or is_skip_header(lines[i-1])):
            # Found a page number marker
            if current_page_num is not None and current_lines:
                pages.append((current_page_num, current_lines))
            
            current_page_num = int(stripped)
            current_lines = []
            i += 1
            # Skip header lines after page number
            skip_count = 0
            while i < len(lines) and skip_count < config['page_header_lines']:
                if is_skip_header(lines[i]):
                    skip_count += 1
                else:
                    break
                i += 1
            continue
        
        if current_page_num is not None or stripped:
            if not is_skip_header(stripped):
                current_lines.append(stripped)
        
        i += 1
    
    # Last page
    if current_page_num is not None and current_lines:
        pages.append((current_page_num, current_lines))
    
    return pages

def fix_text(text):
    """Apply Spanish fixes to OCR text."""
    text = text.replace('\ufffd', '')
    # Fix i -> inverted question mark at start of questions
    text = re.sub(r'\biQu\b', '¿Qu', text)
    text = re.sub(r'\biSois\b', '¿Sois', text)
    text = re.sub(r'\biC(o|u|ómo|ual)\b', lambda m: '¿C' + m.group(1), text)
    text = re.sub(r'\biPor qué\b', '¿Por qué', text)
    text = re.sub(r'\biPor qu\b', '¿Por qu', text)
    text = re.sub(r'\biNo\b', '¿No', text)
    text = re.sub(r'\biLa\b', '¿La', text)
    text = re.sub(r'\biDonde\b', '¿Dónde', text)
    text = re.sub(r'\biEn qué\b', '¿En qué', text)
    text = re.sub(r'\biEn qu\b', '¿En qu', text)
    text = re.sub(r'\biExplicad\b', '¿Explicad', text)
    text = re.sub(r'\biDesde\b', '¿Desde', text)
    text = re.sub(r'\biHab[^a-z]', '¿Hab', text)
    text = re.sub(r'\biA qué\b', '¿A qué', text)
    text = re.sub(r'\biA qu\b', '¿A qu', text)
    text = re.sub(r'\biSu\b', '¿Su', text)
    text = re.sub(r'\biTen[^a-z]', '¿Ten', text)
    text = re.sub(r'\biEsper[^a-z]', '¿Esper', text)
    text = re.sub(r'\biPued[^a-z]', '¿Pued', text)
    text = re.sub(r'\biPrest[^a-z]', '¿Prest', text)
    text = re.sub(r'\biTrabaj[^a-z]', '¿Trabaj', text)
    text = re.sub(r'\biEst[^a-z]', '¿Est', text)
    text = re.sub(r'\biSobre\b', '¿Sobre', text)
    text = re.sub(r'\biOs sug\b', '¿Os sug', text)
    text = re.sub(r'\biQue\b', '¿Que', text)
    text = re.sub(r'\biQuién\b', '¿Quién', text)
    text = re.sub(r'\biQui.n\b', '¿Quién', text)
    text = text.replace(';C', '¿C')
    text = text.replace(';Como', '¿Cómo')
    text = text.replace(';Donde', '¿Dónde')
    # Fix j at start of line -> ¿
    if text.startswith('j'):
        text = '¿' + text[1:]
    text = text.replace('\n;', '\n¿')
    text = text.replace('\ni', '\n¿')
    return text


def classify_line(line):
    if not line.strip():
        return 'empty'
    if 'CATECISMO' in line or line.startswith('INSTRUCCION'):
        return 'header'
    if line.startswith('Pregunta') or line.startswith('¿Pregunta'):
        return 'question'
    if line.startswith('Respuesta'):
        return 'answer'
    if 'A la Gloria' in line or 'A LA GLORIA' in line:
        return 'gloria'
    if 'RITO' in line and 'ANTIGUO' in line:
        return 'rito'
    if 'ORDEN DE LOS RITOS UNIDOS' in line:
        return 'orden'
    if 'GRAN LOGIA' in line:
        return 'granlogia'
    if 'Ritos reunificados' in line or 'Ritos unificados' in line:
        return 'reunificados'
    if line.startswith('PROCLAMACIÓN') or line.startswith('PROCLAMACION'):
        return 'proclamation'
    if line.startswith('CAPÍTULO') or line.startswith('CAPITULO'):
        return 'chapter'
    if line.startswith('TÍTULO') or line.startswith('TITULO'):
        return 'title'
    if line.startswith('Art.') or line.startswith('Artículo') or line.startswith('ARTICULO'):
        return 'article'
    if line.startswith('(§') or line.startswith('(*)') or line.startswith('\u2794') or (line.startswith('(') and line.endswith(')')):
        return 'stage'
    if line.startswith('CHEQUETET') or line.startswith('MAK-BENAH') or line.startswith('ARELICH') or line.startswith('VOMALITES'):
        return 'aclamacion'
    if line.startswith('Amén') or line.startswith('Soberano Maestro') or line.startswith('Eterno Regenerador'):
        return 'oracion'
    if line.startswith('ÍNDICE') or line.startswith('INDICE'):
        return 'header'
    return 'text'


def make_html(config, pages, output_dir):
    lines_out = []
    lines_out.append('<!DOCTYPE html>\n<html lang="es">\n<head>')
    lines_out.append('<meta charset="UTF-8">')
    lines_out.append('<meta name="viewport" content="width=device-width, initial-scale=1.0">')
    lines_out.append(f'<title>{config["title"]}</title>')
    lines_out.append('''<style>
* { margin: 0; padding: 0; box-sizing: border-box; }
body { font-family: 'Times New Roman', Times, serif; background: #f5f0e8; color: #1a1a1a; padding: 40px 20px; }
.container { max-width: 900px; margin: 0 auto; background: #fffdf9; padding: 60px 80px; box-shadow: 0 2px 20px rgba(0,0,0,0.1); border: 1px solid #d4c9b0; }
.page-marker { text-align: center; margin: 40px 0 20px; padding: 10px 0; border-top: 2px solid #8b7355; border-bottom: 2px solid #8b7355; font-size: 13px; font-weight: bold; color: #8b7355; letter-spacing: 3px; }
.page-marker.first { margin-top: 0; }
h1 { text-align: center; font-size: 20px; margin: 30px 0 15px; color: #2c2c2c; }
h2 { text-align: center; font-size: 17px; margin: 25px 0 12px; color: #444; }
h3 { text-align: center; font-size: 15px; margin: 20px 0 10px; color: #555; }
.gloria { font-size: 14px; font-variant: small-caps; letter-spacing: 2px; color: #8b7355; margin: 20px 0; text-align: center; }
.rito { font-size: 16px; font-weight: bold; text-transform: uppercase; letter-spacing: 2px; margin: 15px 0; text-align: center; }
.orden { font-size: 14px; font-style: italic; color: #555; margin: 15px 0; text-align: center; }
.granlogia { font-size: 13px; font-weight: bold; color: #8b7355; margin: 10px 0; text-align: center; }
.reunificados { font-size: 14px; font-style: italic; text-align: center; margin: 10px 0; color: #555; }
.text-block { margin: 8px 0; line-height: 1.6; text-align: justify; }
.proclamation { font-size: 14px; font-weight: bold; text-align: center; margin: 15px 0; letter-spacing: 1px; }
.chapter { font-size: 15px; font-weight: bold; margin: 20px 0 10px; }
.title-header { font-size: 14px; font-weight: bold; font-variant: small-caps; margin: 15px 0 8px; }
.article { margin: 6px 0 6px 10px; }
.instruccion { font-size: 18px; font-weight: bold; text-align: center; margin: 30px 0 20px; }
.catecismo { font-size: 16px; font-weight: bold; text-align: center; margin: 20px 0; font-variant: small-caps; }
.question { margin: 15px 0 5px; font-style: italic; color: #2c2c2c; }
.answer { margin: 0 0 15px 20px; padding-left: 15px; border-left: 3px solid #d4c9b0; color: #333; }
.stage { font-style: italic; color: #666; margin: 10px 0; padding-left: 20px; }
.aclamacion { text-align: center; font-weight: bold; margin: 15px 0; letter-spacing: 1px; }
.oracion { text-align: center; font-style: italic; margin: 15px 0; color: #444; }
@media print { body { background: white; padding: 0; } .container { box-shadow: none; border: none; padding: 40px; } }
</style></head><body><div class="container">''')
    
    # Title
    lines_out.append(f'<h1>{config["title"]}</h1>')
    
    for idx, (page_num, page_lines) in enumerate(pages):
        cls = 'page-marker first' if idx == 0 else 'page-marker'
        lines_out.append(f'<div class="{cls}">— PÁGINA {page_num} —</div>')
        
        for line in page_lines:
            if not line.strip():
                continue
            cls = classify_line(line)
            if cls == 'empty':
                continue
            elif cls == 'header':
                if 'CATECISMO' in line:
                    lines_out.append(f'<div class="catecismo">{line}</div>')
                else:
                    lines_out.append(f'<div class="instruccion">{line}</div>')
            elif cls == 'question':
                q = line.replace('Pregunta', '<strong>Pregunta</strong>')
                lines_out.append(f'<div class="question">{q}</div>')
            elif cls == 'answer':
                a = line.replace('Respuesta', '<strong>Respuesta</strong>')
                lines_out.append(f'<div class="answer">{a}</div>')
            elif cls == 'gloria':
                lines_out.append(f'<div class="gloria">{line}</div>')
            elif cls == 'rito':
                line = line.replace('ANTIGUOYPRIMITIVODEMENFIS', 'ANTIGUO Y PRIMITIVO DE MENFIS')
                lines_out.append(f'<div class="rito">{line}</div>')
            elif cls == 'orden':
                lines_out.append(f'<div class="orden">{line}</div>')
            elif cls == 'granlogia':
                lines_out.append(f'<div class="granlogia">{line}</div>')
            elif cls == 'reunificados':
                lines_out.append(f'<div class="reunificados">{line}</div>')
            elif cls == 'proclamation':
                lines_out.append(f'<div class="proclamation">{line}</div>')
            elif cls == 'chapter':
                lines_out.append(f'<div class="chapter">{line}</div>')
            elif cls == 'title':
                lines_out.append(f'<div class="title-header">{line}</div>')
            elif cls == 'article':
                lines_out.append(f'<div class="article">{line}</div>')
            elif cls == 'stage':
                lines_out.append(f'<div class="stage">{line}</div>')
            elif cls == 'aclamacion':
                lines_out.append(f'<div class="aclamacion">{line}</div>')
            elif cls == 'oracion':
                lines_out.append(f'<div class="oracion">{line}</div>')
            else:
                lines_out.append(f'<div class="text-block">{line}</div>')
    
    lines_out.append('</div></body></html>')
    
    fname = f'{config["id"]}.html'
    path = os.path.join(output_dir, fname)
    with open(path, 'w', encoding='utf-8') as f:
        f.write('\n'.join(lines_out))
    print(f'  HTML: {fname}')


def make_md(config, pages, output_dir):
    lines_out = []
    lines_out.append(f'# {config["title"]}')
    lines_out.append('')
    
    for idx, (page_num, page_lines) in enumerate(pages):
        lines_out.append('')
        lines_out.append(f'--- **PÁGINA {page_num}** ---')
        lines_out.append('')
        
        for line in page_lines:
            if not line.strip():
                continue
            cls = classify_line(line)
            if cls == 'empty':
                continue
            elif cls == 'header':
                if 'CATECISMO' in line:
                    lines_out.append(f'## {line}')
                else:
                    lines_out.append(f'### {line}')
            elif cls == 'question':
                q = line.replace('Pregunta', '**Pregunta:**')
                lines_out.append(q)
            elif cls == 'answer':
                a = line.replace('Respuesta', '**Respuesta:**')
                lines_out.append(f'> {a}')
            elif cls == 'gloria':
                lines_out.append(f'*{line}*')
            elif cls == 'rito':
                line = line.replace('ANTIGUOYPRIMITIVODEMENFIS', 'ANTIGUO Y PRIMITIVO DE MENFIS')
                lines_out.append(f'# {line}')
            elif cls in ['orden', 'reunificados']:
                lines_out.append(f'*{line}*')
            elif cls == 'granlogia':
                lines_out.append(f'**{line}**')
            elif cls == 'proclamation':
                lines_out.append(f'## {line}')
            elif cls == 'chapter':
                lines_out.append(f'## {line}')
            elif cls == 'title':
                lines_out.append(f'### {line}')
            elif cls in ['article', 'text']:
                lines_out.append(line)
            elif cls == 'stage':
                lines_out.append(f'*{line}*')
            elif cls in ['aclamacion', 'oracion']:
                lines_out.append(f'**{line}**' if cls == 'aclamacion' else f'*{line}*')
            lines_out.append('')
    
    fname = f'{config["id"]}.md'
    path = os.path.join(output_dir, fname)
    with open(path, 'w', encoding='utf-8') as f:
        f.write('\n'.join(lines_out))
    print(f'  MD:   {fname}')


def add_docx_paragraph(doc, text, style_name='Normal', bold=False, italic=False, alignment=None,
                       font_size=None, indent=None, font_name='Times New Roman'):
    p = doc.add_paragraph()
    if alignment:
        p.alignment = alignment
    if indent:
        p.paragraph_format.left_indent = indent
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = font_name
    if font_size:
        run.font.size = font_size
    return p


def make_docx(config, pages, output_dir):
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(11)
    
    # Title
    add_docx_paragraph(doc, config['title'], bold=True, font_size=Pt(18),
                       alignment=WD_ALIGN_PARAGRAPH.CENTER)
    
    for idx, (page_num, page_lines) in enumerate(pages):
        # Page marker
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f'—— PÁGINA {page_num} ——')
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x8B, 0x73, 0x55)
        
        for line in page_lines:
            if not line.strip():
                continue
            cls = classify_line(line)
            if cls == 'empty':
                continue
            
            if cls == 'header':
                is_catecismo = 'CATECISMO' in line
                add_docx_paragraph(doc, line, bold=True,
                                   font_size=Pt(14) if not is_catecismo else Pt(13),
                                   alignment=WD_ALIGN_PARAGRAPH.CENTER)
            elif cls == 'question':
                q = line.replace('Pregunta', 'Pregunta:')
                add_docx_paragraph(doc, q, italic=True)
            elif cls == 'answer':
                a = line.replace('Respuesta', 'Respuesta:')
                add_docx_paragraph(doc, a, indent=Inches(0.4))
            elif cls == 'gloria':
                add_docx_paragraph(doc, line, italic=True, font_size=Pt(11),
                                   alignment=WD_ALIGN_PARAGRAPH.CENTER)
            elif cls == 'rito':
                line = line.replace('ANTIGUOYPRIMITIVODEMENFIS', 'ANTIGUO Y PRIMITIVO DE MENFIS')
                add_docx_paragraph(doc, line, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
            elif cls == 'orden' or cls == 'reunificados':
                add_docx_paragraph(doc, line, italic=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
            elif cls == 'granlogia':
                add_docx_paragraph(doc, line, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
            elif cls == 'proclamation':
                add_docx_paragraph(doc, line, bold=True, font_size=Pt(13),
                                   alignment=WD_ALIGN_PARAGRAPH.CENTER)
            elif cls == 'chapter':
                add_docx_paragraph(doc, line, bold=True, font_size=Pt(12))
            elif cls == 'title':
                add_docx_paragraph(doc, line, bold=True)
            elif cls in ['article', 'text']:
                add_docx_paragraph(doc, line)
            elif cls == 'stage':
                add_docx_paragraph(doc, line, italic=True)
            elif cls == 'aclamacion':
                add_docx_paragraph(doc, line, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
            elif cls == 'oracion':
                add_docx_paragraph(doc, line, italic=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    
    fname = f'{config["id"]}.docx'
    path = os.path.join(output_dir, fname)
    doc.save(path)
    print(f'  DOCX: {fname}')


def process_source(config):
    print(f'\n{config["title"]}')
    text = load_text(config['file'])
    text = fix_text(text)
    pages = split_into_pages(text, config)
    print(f'  Páginas encontradas: {len(pages)}')
    if pages:
        pgnums = [p[0] for p in pages]
        print(f'  Números: {pgnums[:5]}...{pgnums[-3:]}')
    
    html_dir = os.path.join(OUTPUT, 'html')
    md_dir = os.path.join(OUTPUT, 'md')
    doc_dir = os.path.join(OUTPUT, 'doc')
    for d in [html_dir, md_dir, doc_dir]:
        os.makedirs(d, exist_ok=True)
    
    make_html(config, pages, html_dir)
    make_md(config, pages, md_dir)
    make_docx(config, pages, doc_dir)


# Run all
print("Generando rituales paginados...")
for src in SOURCES:
    process_source(src)
print("\n¡Todos los rituales generados con paginación!")
