import os, re
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

BASE = os.path.join(os.environ['TEMP'], 'pdf_text_new')
OUTPUT = r'C:\Temp\kitian\GITHUT\RItoMemphisMisraim\Rituales\output'

HTML_TEMPLATE = '''<!DOCTYPE html>
<html lang="es">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{title}</title>
    <style>
        * {{ margin: 0; padding: 0; box-sizing: border-box; }}
        body {{ font-family: 'Times New Roman', Times, serif; background: #f5f0e8; color: #1a1a1a; padding: 40px 20px; }}
        .container {{ max-width: 900px; margin: 0 auto; background: #fffdf9; padding: 60px 80px; box-shadow: 0 2px 20px rgba(0,0,0,0.1); border: 1px solid #d4c9b0; }}
        h1 {{ text-align: center; font-size: 20px; margin: 30px 0 15px; color: #2c2c2c; }}
        h2 {{ text-align: center; font-size: 17px; margin: 25px 0 12px; color: #444; }}
        h3 {{ font-size: 15px; margin: 20px 0 10px; color: #555; }}
        .gloria {{ font-size: 14px; font-variant: small-caps; letter-spacing: 2px; color: #8b7355; margin: 20px 0; text-align: center; }}
        .rito {{ font-size: 16px; font-weight: bold; text-transform: uppercase; letter-spacing: 2px; margin: 15px 0; text-align: center; }}
        .gran-logia {{ font-size: 13px; font-weight: bold; color: #8b7355; margin: 10px 0; text-align: center; }}
        .text-block {{ margin: 8px 0; line-height: 1.6; text-align: justify; }}
        .cita {{ margin: 10px 20px; padding-left: 15px; border-left: 3px solid #d4c9b0; color: #444; font-style: italic; }}
        .aclamacion {{ text-align: center; font-weight: bold; margin: 15px 0; letter-spacing: 1px; }}
        .oracion {{ text-align: center; font-style: italic; margin: 15px 0; }}
        .stage {{ font-style: italic; color: #666; margin: 10px 0; padding-left: 20px; }}
        .footer {{ text-align: center; margin-top: 40px; padding-top: 20px; border-top: 1px solid #d4c9b0; font-size: 12px; color: #888; }}
        @media print {{ body {{ background: white; padding: 0; }} .container {{ box-shadow: none; border: none; padding: 40px; }} }}
    </style>
</head>
<body>
    <div class="container">
        {content}
    </div>
</body>
</html>'''

def clean_line(line):
    line = line.strip()
    # Remove page numbers
    if re.match(r'^\d+$', line):
        return ''
    # Remove orphan header fragments
    if line in ['S O', '. G', 'S CO S', 'UC O', 'S', 'G', 'OS G', 'S.']:
        return ''
    return line

def read_and_split(path):
    with open(path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    cleaned = []
    for line in lines:
        c = clean_line(line)
        if c:
            cleaned.append(c)
    return cleaned

def classify_line(line):
    """Classify a line for styling"""
    if 'A la Gloria' in line or 'A LA GLORIA' in line:
        return 'gloria'
    if 'RITO ANTIGUO' in line or 'RITO ANTIGUO' in line.upper():
        return 'rito'
    if 'GRAN LOGIA DE ESPANA' in line or 'GRAN LOGIA DE ESPA' in line:
        return 'gran-logia'
    if line.startswith('GRANDES CONSTITUCIONES') or line.startswith('REGLAMENTOS GENERALES'):
        return 'header1'
    if line.startswith('RITUALES, CEREMONIAS'):
        return 'header2'
    if line.startswith('INSTRUCCION') or line.startswith('CATECISMO'):
        return 'header3'
    if line.startswith('CAPITULO') or line.startswith('CAP\u00cdTULO') or line.startswith('TITULO') or line.startswith('T\u00cdTULO'):
        return 'header1'
    if line.startswith('Pregunta') or line.startswith('\u00bfPregunta'):
        return 'question'
    if line.startswith('Respuesta'):
        return 'answer'
    if line.startswith('CHEQUETET') or line.startswith('MAK-BENAH'):
        return 'aclamacion'
    if line.startswith('Am\u00e9n') or line.startswith('Soberano Maestro'):
        return 'oracion'
    if line.startswith('(') or line.startswith('\u2794') or line.startswith('* '):
        return 'stage'
    if line.startswith('\u00bf') or line.startswith(';') or line.startswith('i'):
        return 'question_text'
    if line.startswith('\u2014') or line.startswith('-'):
        return 'cita'
    if line.startswith('\u00cdNDICE') or line.startswith('INDICE'):
        return 'header1'
    return 'text'

def build_html(lines, title):
    parts = []
    parts.append(f'<h1>{title}</h1>')
    
    for line in lines:
        cls = classify_line(line)
        if cls == 'gloria':
            parts.append(f'<div class="gloria">{line}</div>')
        elif cls == 'rito':
            parts.append(f'<div class="rito">{line}</div>')
        elif cls == 'gran-logia':
            parts.append(f'<div class="gran-logia">{line}</div>')
        elif cls == 'header1':
            parts.append(f'<h2>{line}</h2>')
        elif cls == 'header2':
            parts.append(f'<div class="text-block" style="text-align:center;font-weight:bold;font-size:14px;">{line}</div>')
        elif cls == 'header3':
            parts.append(f'<h3>{line}</h3>')
        elif cls == 'question':
            parts.append(f'<div class="text-block" style="font-style:italic;font-weight:bold;margin-top:12px;">{line}</div>')
        elif cls == 'answer':
            parts.append(f'<div class="text-block" style="margin-left:15px;padding-left:10px;border-left:2px solid #d4c9b0;">{line}</div>')
        elif cls == 'aclamacion':
            parts.append(f'<div class="aclamacion">{line}</div>')
        elif cls == 'oracion':
            parts.append(f'<div class="oracion">{line}</div>')
        elif cls == 'stage':
            parts.append(f'<div class="stage">{line}</div>')
        elif cls == 'cita':
            parts.append(f'<div class="cita">{line}</div>')
        elif cls == 'question_text':
            parts.append(f'<div class="text-block" style="font-style:italic;">{line}</div>')
        else:
            parts.append(f'<div class="text-block">{line}</div>')
    
    return '\n'.join(parts)

def build_md(lines, title):
    parts = []
    parts.append(f'# {title}')
    parts.append('')
    
    for line in lines:
        cls = classify_line(line)
        if cls == 'gloria':
            parts.append(f'*{line}*')
        elif cls == 'rito':
            parts.append(f'**{line}**')
        elif cls == 'gran-logia':
            parts.append(f'**{line}**')
        elif cls == 'header1':
            parts.append(f'## {line}')
        elif cls == 'header2':
            parts.append(f'### {line}')
        elif cls == 'header3':
            parts.append(f'#### {line}')
        elif cls == 'question':
            parts.append(f'*{line}*')
        elif cls == 'answer':
            parts.append(f'> {line}')
        elif cls == 'aclamacion':
            parts.append(f'**{line}**')
        elif cls == 'oracion':
            parts.append(f'*{line}*')
        elif cls == 'stage':
            parts.append(f'*{line}*')
        elif cls == 'cita':
            parts.append(f'> {line}')
        elif cls == 'question_text':
            parts.append(f'*{line}*')
        else:
            parts.append(line)
        parts.append('')
    
    return '\n'.join(parts)

def build_docx(lines, title, doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(18)
    
    for line in lines:
        cls = classify_line(line)
        p = doc.add_paragraph()
        
        if cls == 'gloria':
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(line)
            run.italic = True
            run.font.size = Pt(11)
        elif cls == 'rito':
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(line)
            run.bold = True
        elif cls == 'gran-logia':
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(line)
            run.bold = True
        elif cls in ['header1', 'header2']:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(line)
            run.bold = True
            run.font.size = Pt(14)
        elif cls == 'header3':
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(line)
            run.bold = True
        elif cls == 'question':
            run = p.add_run(line)
            run.italic = True
        elif cls == 'answer':
            p.paragraph_format.left_indent = Inches(0.3)
            p.add_run(line)
        elif cls == 'aclamacion':
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(line)
            run.bold = True
        elif cls in ['oracion', 'stage']:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if cls == 'oracion' else WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(line)
            run.italic = True
        elif cls == 'cita':
            p.paragraph_format.left_indent = Inches(0.3)
            run = p.add_run(line)
            run.italic = True
        else:
            p.add_run(line)

def generate_all(name, title, lines):
    html_dir = os.path.join(OUTPUT, 'html')
    md_dir = os.path.join(OUTPUT, 'md')
    doc_dir = os.path.join(OUTPUT, 'doc')
    os.makedirs(html_dir, exist_ok=True)
    os.makedirs(md_dir, exist_ok=True)
    os.makedirs(doc_dir, exist_ok=True)
    
    # HTML
    html_content = build_html(lines, title)
    html = HTML_TEMPLATE.format(title=title, content=html_content)
    html_path = os.path.join(html_dir, f'{name}.html')
    with open(html_path, 'w', encoding='utf-8') as f:
        f.write(html)
    print(f'  HTML: {name}.html')
    
    # MD
    md_content = build_md(lines, title)
    md_path = os.path.join(md_dir, f'{name}.md')
    with open(md_path, 'w', encoding='utf-8') as f:
        f.write(md_content)
    print(f'  MD:   {name}.md')
    
    # DOCX
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(11)
    build_docx(lines, title, doc)
    docx_path = os.path.join(doc_dir, f'{name}.docx')
    doc.save(docx_path)
    print(f'  DOCX: {name}.docx')

print("Generando nuevos documentos...\n")

# 1. Grandes Constituciones y Reglamentos Generales
print("1. Grandes Constituciones y Reglamentos Generales")
lines_0 = read_and_split(os.path.join(BASE, 'ritual_0_completo.txt'))
generate_all('Grandes_Constituciones_y_Reglamentos_Generales',
             'Grandes Constituciones y Reglamentos Generales', lines_0)

# 2. Rituales Simbólicos Completo
print("\n2. Rituales Simbólicos - Obra Completa")
lines_4 = read_and_split(os.path.join(BASE, 'ritual_4_completo.txt'))
generate_all('Rituales_Simbolicos_Completo',
             'Rituales e Instrucciones de los Grados Simb\u00f3licos - Obra Completa', lines_4)

print("\n\u00a1Generados exitosamente!")
