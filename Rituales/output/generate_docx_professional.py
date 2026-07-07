import fitz, os, re, sys
from PIL import Image
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT = r'C:\Temp\kitian\GITHUT\RItoMemphisMisraim\Rituales\output'
SOURCE = os.path.join(os.environ['TEMP'], 'pdf_text_new', 'ritual_4_completo.txt')
PDF_PATH = r'C:\Temp\kitian\GITHUT\RItoMemphisMisraim\Rituales\completo Ritual.pdf'
IMG_DIR = os.path.join(OUTPUT, 'images')

# Words to skip from bolding
SKIP_BOLD = {'Y','E','O','A','DE','DEL','EN','LA','EL','LOS','LAS','UN','UNA',
             'POR','CON','QUE','ES','NO','SE','SU','AL','LO','LE','SUS','HA',
             'HE','HOY','VA','DA','SAN','ERA','SON','HAN','SER','MAS',
             'I','III','II','IV','VI','VII','VIII','IX','XI','XII','XIII',
             '1ER','2O','3ER','D','L','M','H','V','VV','MM','HH',
             'AAP','CC','TT','S','A','N','O','P','R','B','C','J','K',
             'G','T','HO','H','S','N','O','R','F','P','J','B','Z'}


# Speaker patterns for ritual dialog
SPEAKER_SET = {
    'V∴M∴','V∴M','V∴ M∴','V∴ M','M∴ de C∴','M∴ de  C∴','M∴D∴C∴','M∴D∴C∴I∴',
    'V∴ M∴de C∴','V∴ M∴ de C∴',
    '1er Vig∴','1er  Vig∴','E∴V∴','2º Vig∴','2º  Vig∴','2º Vig∴ Electo',
    '1er Vig∴ Electo','V∴1er Vig∴','V∴2º Vig∴',
    'Orad∴','Or∴','V∴Orador','V∴Orador :',
    'Sec∴','V∴Secretario','Tes∴','Hos∴','Exp∴','Experto∴',
    'Cand∴','G∴T∴','V∴G∴T∴','V∴Hierofante',
    'Col∴de Arm∴','Col∴ de Ar∴','Col∴ de  Ar∴',
    '1er Inst∴','2º Inst∴','Expert Inst∴','Experto Inst∴',
    'Los 2 VVig∴','TT∴ los OOf∴','MM∴MM∴',
    'V∴M∴ Electo','V∴M ∴','V∴M∴ :','M∴ de C∴:',
    'Todos los HH∴',
}

SPEAKER_RE = re.compile(r'^(' + '|'.join(re.escape(s) for s in sorted(SPEAKER_SET, key=len, reverse=True)) + r')\s*[:\u2794]?\s*$')

def is_speaker_line(s):
    s = s.strip()
    if not s:
        return False
    return bool(SPEAKER_RE.match(s))

def group_dialog_lines(content_lines):
    segments = []
    i = 0
    while i < len(content_lines):
        s = content_lines[i].strip()
        if is_speaker_line(s):
            speaker = s
            speech_lines = []
            i += 1
            while i < len(content_lines):
                ns = content_lines[i].strip()
                if not ns:
                    i += 1
                    continue
                if is_speaker_line(ns):
                    break
                speech_lines.append(ns)
                i += 1
            segments.append(('dialog', speaker, speech_lines))
        else:
            segments.append(('line', s))
            i += 1
    return segments

def has_uppercase_word(word):
    word_clean = word.strip('.,;:!?¿¡"\'()[]{}«»-–—')
    if not word_clean or len(word_clean) < 3:
        return False
    if word_clean in SKIP_BOLD:
        return False
    alpha_chars = [c for c in word_clean if c.isalpha()]
    if not alpha_chars:
        return False
    return all(c.isupper() for c in alpha_chars)

def load_source():
    with open(SOURCE, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Reemplazar el nombre de la Gran Logia usando regex para manejar saltos de línea
    import re
    
    # Patrón principal: maneja saltos de línea y espacios variables
    pattern = r'Gran\s+Logia\s+de\s+Espa[ñn]a\s+de\s+Menfis[-\s]+Mizraim\s+del\s+Rito\s+Antiguo\s+y\s+Primitivo\s+de\s+Menfis[-\s]+Mizraim\s+y\s+de\s+la\s+Orden\s+de\s+los\s+Ritos\s+Unidos\s+de\s+Menfis\s*[&y]\s*Mizraim'
    new_name = 'Gran Logia Simbólica del Rito Antiguo y Primitivo de Menfis Misraim de la República del Paraguay'
    content = re.sub(pattern, new_name, content, flags=re.IGNORECASE)
    
    # Variante con "de la del"
    pattern2 = r'Gran\s+Logia\s+de\s+Espa[ñn]a\s+de\s+Menfis[-\s]+Mizraim\s+de\s+la\s+del\s+Rito\s+Antiguo\s+y\s+Primitivo\s+de\s+Menfis[-\s]+Mizraim\s+y\s+de\s+la\s+Orden\s+de\s+los\s+Ritos\s+Unidos\s+de\s+Menfis\s*[&y]\s*Mizraim'
    content = re.sub(pattern2, new_name, content, flags=re.IGNORECASE)
    
    # Variante más corta (solo "Gran Logia de España de Menfis-Mizraim" seguida de "del Rito")
    pattern3 = r'Gran\s+Logia\s+de\s+Espa[ñn]a\s+de\s+Menfis[-\s]+Mizraim\s+del\s+Rito\s+Antiguo\s+y\s+Primitivo'
    content = re.sub(pattern3, 'Gran Logia Simbólica del Rito Antiguo y Primitivo', content, flags=re.IGNORECASE)
    
    # Variante más corta aún
    pattern4 = r'Gran\s+Logia\s+de\s+Espa[ñn]a\s+de\s+Menfis[-\s]+Mizraim'
    content = re.sub(pattern4, 'Gran Logia Simbólica', content, flags=re.IGNORECASE)
    
    # Eliminar la referencia a "Ritos unificados por José GARIBALDI en 1881"
    pattern5 = r'Gran\s+Logia\s+Simbólica\s+del\s+Rito\s+Antiguo\s+y\s+Primitivo\s+de\s+Menfis\s+Misraim\s+de\s+la\s+República\s+del\s+Paraguay\s*[–\-—]\s*Ritos\s+(?:unificados|reunificados)\s+por\s+Jos[ée]\s+GARIBALDI\s+en\s+1881'
    content = re.sub(pattern5, 'Gran Logia Simbólica del Rito Antiguo y Primitivo de Menfis Misraim de la República del Paraguay', content, flags=re.IGNORECASE)
    
    # Reemplazar "en el Oriente de:………………..." por "en el Oriente del Paraguay"
    pattern6 = r'en\s+el\s+Oriente\s+de\s*[:\.]+\s*'
    content = re.sub(pattern6, 'en el Oriente del Paraguay ', content, flags=re.IGNORECASE)
    
    # Unir líneas fragmentadas innecesariamente
    # Detectar líneas que terminan sin puntuación final y unirlas con la siguiente
    lines = content.split('\n')
    joined_lines = []
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        # Si la línea no termina con puntuación final ni está vacía, unirla con la siguiente
        if line and not line.endswith(('.', '!', '?', ':', '…', '∴')) and not line.startswith('---') and i + 1 < len(lines):
            next_line = lines[i + 1].strip()
            # Unir si la siguiente línea no es un marcador de página ni está vacía
            if next_line and not next_line.startswith('--- PÁGINA') and not next_line.startswith('--- **PÁGINA'):
                joined_lines.append(line + ' ' + next_line)
                i += 2
                continue
        joined_lines.append(line)
        i += 1
    
    content = '\n'.join(joined_lines)
    
    # Limpiar secuencia redundante en portadas/encabezados
    # Patrón: Gran Logia Simbólica 2011 + RITO ANTIGUO... + ORDEN DE LOS RITOS... + Ritos reunificados... + Gran Logia Simbólica
    # Reemplazar por solo: RITO ANTIGUO Y PRIMITIVO DE MENFIS-MIZRAIM
    pattern7 = r'Gran\s+Logia\s+Simbólica\s+2011\s+RITO\s+ANTIGUO\s+Y\s+PRIMITIVO\s+DE\s+MENFIS-MIZRAIM\s+ORDEN\s+DE\s+LOS\s+RITOS\s+UNIDOS\s+DE\s+MENFIS\s*[&y]\s*MIZRAIM\s+Ritos\s+(?:reunificados|unificados)\s+por\s+Jos[ée]\s+GARIBALDI\s+en\s+1881\s+Gran\s+Logia\s+Simbólica'
    content = re.sub(pattern7, 'RITO ANTIGUO Y PRIMITIVO DE MENFIS-MIZRAIM', content, flags=re.IGNORECASE)
    
    # Eliminar "Gran Logia Simbólica 2011" (línea suelta)
    pattern8 = r'Gran\s+Logia\s+Simbólica\s+2011\s*\n?'
    content = re.sub(pattern8, '', content, flags=re.IGNORECASE)
    
    return content

def extract_pages(text):
    lines = text.split('\n')
    line_freq = {}
    for line in lines:
        stripped = line.strip()
        if stripped and not stripped.isdigit():
            line_freq[stripped] = line_freq.get(stripped, 0) + 1
    recurring = {l for l, c in line_freq.items() if c >= 3}

    skip_prefixes = [
        'GRAN LOGIA DE ESPANA DE MENFIS-MIZRAIM',
        'GRAN LOGIA DE ESPAÑA DE MENFIS-MIZRAIM',
        'A LA GLORIA DEL SUBLIME ARQUITECTO DE LOS MUNDOS',
        'A la Gloria del Sublime Arquitecto de los Mundos',
        'Bajo la Autoridad Espiritual',
        'Federacion de Logias',
        'Federación de Logias',
        'Traducidos del original',
        'RITUALES, CEREMONIAS E INSTRUCCIONES',
        'Rituales e Instrucciones de los',
        'Rituales, Ceremonias e Instrucciones',
        'HISTORIA. GRANDES CONSTITUCIONES',
        'HISTORIA.GRANDES CONSTITUCIONES',
    ]

    pages = {}
    current_page = None
    current_lines = []

    for line in lines:
        stripped = line.strip()
        if stripped.isdigit() and 2 <= int(stripped) <= 400:
            num = int(stripped)
            if current_page is not None and current_lines:
                pages[current_page] = current_lines
            current_page = num
            current_lines = []
            continue
        if current_page is not None:
            if not stripped:
                continue
            if stripped in recurring:
                skip_it = False
                for pref in skip_prefixes:
                    if stripped.startswith(pref):
                        skip_it = True
                        break
                if skip_it:
                    continue
            if re.match(r'^[\s\u00b7\u2022\u2023\u25cf\u25d8\u2713\u2714\u2666\u2663\u2605\u2606\u2234\*\-\.\~\,\s]+$', stripped):
                continue
            current_lines.append(stripped)

    if current_page is not None and current_lines:
        pages[current_page] = current_lines
    return pages

def add_page_number(footer):
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    run1 = p.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    run1._r.append(fldChar1)
    
    run2 = p.add_run()
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = ' PAGE '
    run2._r.append(instrText)
    
    run3 = p.add_run()
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run3._r.append(fldChar2)
    
    for r in [run1, run2, run3]:
        r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(0x8B, 0x73, 0x55)

def add_header_to_section(section, logo_path=None):
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    if logo_path and os.path.exists(logo_path):
        run = hp.add_run()
        run.add_picture(logo_path, width=Inches(0.5))
        run2 = hp.add_run('\n')
        run2.font.size = Pt(8)
    
    run3 = hp.add_run('RITUALES, CEREMONIAS E INSTRUCCIONES DE LOS GRADOS SIMBÓLICOS')
    run3.font.size = Pt(7)
    run3.font.name = 'Times New Roman'
    run3.font.color.rgb = RGBColor(0x8B, 0x73, 0x55)
    
    run4 = hp.add_run('\nHISTORIA. GRANDES CONSTITUCIONES Y REGLAMENTOS GENERALES.')
    run4.font.size = Pt(7)
    run4.font.name = 'Times New Roman'
    run4.font.color.rgb = RGBColor(0x8B, 0x73, 0x55)
    
    run5 = hp.add_run('\nGRAN LOGIA DE ESPAÑA DE MENFIS-MIZRAIM 2011')
    run5.font.size = Pt(7)
    run5.font.name = 'Times New Roman'
    run5.font.color.rgb = RGBColor(0x8B, 0x73, 0x55)

def add_para(doc, text, bold=False, italic=False, alignment=None, font_size=None, indent=None, auto_bold=True, heading_level=0, space_after=None, space_before=None):
    if heading_level:
        p = doc.add_heading(text, level=heading_level)
        if alignment:
            p.alignment = alignment
        if space_after is not None:
            p.paragraph_format.space_after = space_after
        if space_before is not None:
            p.paragraph_format.space_before = space_before
        return p

    p = doc.add_paragraph()
    if alignment:
        p.alignment = alignment
    if indent:
        p.paragraph_format.left_indent = indent
    if space_after is not None:
        p.paragraph_format.space_after = space_after
    if space_before is not None:
        p.paragraph_format.space_before = space_before

    if auto_bold and not bold:
        words = text.split()
        for i, w in enumerate(words):
            if has_uppercase_word(w):
                prefix = ''
                suffix = ''
                clean = w
                while clean and not clean[-1].isalpha():
                    suffix = clean[-1] + suffix
                    clean = clean[:-1]
                while clean and not clean[0].isalpha():
                    prefix = prefix + clean[0]
                    clean = clean[1:]
                if prefix:
                    r = p.add_run(prefix)
                    r.font.name = 'Times New Roman'
                    if italic: r.italic = True
                r = p.add_run(clean)
                r.bold = True
                r.font.name = 'Times New Roman'
                if italic: r.italic = True
                if font_size: r.font.size = font_size
                if suffix:
                    r = p.add_run(suffix)
                    r.font.name = 'Times New Roman'
                    if italic: r.italic = True
            else:
                r = p.add_run(w)
                r.font.name = 'Times New Roman'
                if italic: r.italic = True
                if font_size: r.font.size = font_size
            if i < len(words) - 1:
                r = p.add_run(' ')
                r.font.name = 'Times New Roman'
    else:
        run = p.add_run(text)
        run.bold = bold
        run.italic = italic
        run.font.name = 'Times New Roman'
        if font_size:
            run.font.size = font_size
    return p

def generate_docx(pages_dict, doc_config):
    page_nums = doc_config['pages']
    title = doc_config['title']
    is_completo = doc_config['id'] == 'Rituales_Simbolicos_Completo'

    doc = Document()
    
    # ========================
    # PAGE SETUP: A4, margins
    # ========================
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.0)
    section.different_first_page_header_footer = True
    
    # ========================
    # NORMAL STYLE
    # ========================
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(10)
    style.paragraph_format.line_spacing = 1.15
    style.paragraph_format.space_after = Pt(2)
    style.paragraph_format.space_before = Pt(0)
    
    # ========================
    # HEADING STYLES
    # ========================
    for level in range(1, 5):
        hs = doc.styles[f'Heading {level}']
        hs.font.name = 'Times New Roman'
        hs.font.color.rgb = RGBColor(0x1a, 0x1a, 0x1a)
        hs.font.bold = True
        if level == 1:
            hs.font.size = Pt(16)
            hs.paragraph_format.space_before = Pt(12)
            hs.paragraph_format.space_after = Pt(6)
        elif level == 2:
            hs.font.size = Pt(13)
            hs.paragraph_format.space_before = Pt(10)
            hs.paragraph_format.space_after = Pt(4)
        elif level == 3:
            hs.font.size = Pt(11)
            hs.paragraph_format.space_before = Pt(8)
            hs.paragraph_format.space_after = Pt(3)
    
    # ========================
    # FOOTER: Page numbers
    # ========================
    add_page_number(section.footer)
    
    # First-page footer (empty)
    first_footer = section.first_page_footer
    first_footer.is_linked_to_previous = False
    
    # ========================
    # HEADER for content pages
    # ========================
    logo_path = os.path.join(IMG_DIR, 'header_logo_0.png')
    add_header_to_section(section, logo_path)
    
    # First-page header (empty - cover doesn't need header)
    first_header = section.first_page_header
    first_header.is_linked_to_previous = False
    
    # ========================
    # COVER PAGE (page 1 from PDF) - only for complete version
    # ========================
    if is_completo:
        cover_img = os.path.join(IMG_DIR, 'cover.png')
        if os.path.exists(cover_img):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(cover_img, width=Cm(17.0))
        
        doc.add_page_break()
        
        # PAGE 4: ÍNDICE (TOC) - render as image (only for complete)
        toc_img = os.path.join(IMG_DIR, 'page4_img_2.png')
        if os.path.exists(toc_img):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(toc_img, width=Cm(17.0))
        
        doc.add_page_break()
    
    # ========================
    # CONTENT PAGES
    # ========================
    # Diagram pages (463x151 images = esquema de logia)
    diagram_pages = {7, 12, 13, 14, 15, 16, 19, 20, 22, 26, 29, 32, 36, 38, 44, 52, 53, 55, 57, 58, 59, 60, 61, 62, 
                     77, 79, 80, 81, 90, 97, 102, 103, 104, 106, 107, 110, 128, 129, 131, 139, 140, 142, 147, 149, 151,
                     165, 168, 179, 181, 186, 187, 188, 196, 198, 199, 200, 201, 202, 204, 205, 206, 207, 208, 219, 223,
                     232, 257, 258, 259, 260, 261, 263, 277, 278, 279, 280, 281}
    
    # Only skip early pages for the complete version
    start_page = 5 if is_completo else 2
    
    for pg in page_nums:
        if pg < start_page:
            continue
        
        if pg not in pages_dict:
            continue
        
        doc.add_page_break()
        
        # Add diagram image if present
        if pg in diagram_pages:
            # Extract the diagram image from PDF
            pdf_doc = fitz.open(PDF_PATH)
            page_obj = pdf_doc[pg - 1]
            images = page_obj.get_images(full=True)
            for img in images:
                xref = img[0]
                pix = fitz.Pixmap(pdf_doc, xref)
                if pix.width == 463 and pix.height == 151:
                    if pix.n > 4:
                        pix = fitz.Pixmap(fitz.csRGB, pix)
                    diag_path = os.path.join(IMG_DIR, f'diagram_p{pg}.png')
                    pix.save(diag_path)
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = p.add_run()
                    run.add_picture(diag_path, width=Cm(14.0))
                    break
            pdf_doc.close()
        
                # Process content with dialog grouping
        segments = group_dialog_lines(pages_dict[pg])
        for seg in segments:
            if seg[0] == 'dialog':
                speaker_text = seg[1]
                speech_list = seg[2]
                table = doc.add_table(rows=1, cols=2)
                for row in table.rows:
                    for cell in row.cells:
                        tc = cell._tc
                        tcPr = tc.get_or_add_tcPr()
                        tcBorders = OxmlElement('w:tcBorders')
                        for bn in ['top', 'left', 'bottom', 'right']:
                            b = OxmlElement(f'w:{bn}')
                            b.set(qn('w:val'), 'nil')
                            tcBorders.append(b)
                        tcPr.append(tcBorders)
                cell_l = table.cell(0, 0)
                cell_l.width = Cm(3.5)
                cp = cell_l.paragraphs[0]
                run = cp.add_run(speaker_text)
                run.bold = True
                run.font.name = 'Times New Roman'
                run.font.size = Pt(10)
                cell_r = table.cell(0, 1)
                for si, sl in enumerate(speech_list):
                    if si > 0:
                        cell_r.add_paragraph()
                    cp = cell_r.paragraphs[-1]
                    words = sl.split()
                    for wi, w in enumerate(words):
                        if has_uppercase_word(w):
                            prefix = ''
                            suffix = ''
                            clean = w
                            while clean and not clean[-1].isalpha():
                                suffix = clean[-1] + suffix
                                clean = clean[:-1]
                            while clean and not clean[0].isalpha():
                                prefix = prefix + clean[0]
                                clean = clean[1:]
                            if prefix:
                                r = cp.add_run(prefix)
                                r.font.name = 'Times New Roman'
                                r.font.size = Pt(10)
                            r = cp.add_run(clean)
                            r.bold = True
                            r.font.name = 'Times New Roman'
                            r.font.size = Pt(10)
                            if suffix:
                                r = cp.add_run(suffix)
                                r.font.name = 'Times New Roman'
                                r.font.size = Pt(10)
                        else:
                            r = cp.add_run(w)
                            r.font.name = 'Times New Roman'
                            r.font.size = Pt(10)
                        if wi < len(words) - 1:
                            r = cp.add_run(' ')
                            r.font.name = 'Times New Roman'
            else:
                s = seg[1]

                # Heading level detection for TOC
                h = 0
                if s in ['CATECISMO'] or s.startswith('INSTRUCCIÓN DEL') or s.startswith('INSTRUCCION DEL'):
                    h = 2
                elif s == 'ÍNDICE' or s == 'INDICE':
                    h = 1
                elif s.startswith('PROCLAMACIÓN') or s.startswith('PROCLAMACION'):
                    h = 1
                elif s.startswith('CAPÍTULO') or s.startswith('CAPITULO') or s.startswith('CAP.'):
                    h = 2
                elif s == 'TÍTULOS MASÓNICOS' or s == 'TITULOS MASONICOS':
                    h = 2
                elif s.startswith('TÍTULO') or s.startswith('TITULO'):
                    h = 3
                elif s.startswith('DISPOSICIONES'):
                    h = 2
                elif s.startswith('GRAN LOGIA DE ESPAÑA') or s.startswith('GRAN LOGIA DE ESPANA'):
                    h = 1
                elif 'RITO ANTIGUO' in s:
                    h = 1

                if h:
                    add_para(doc, s, bold=True, heading_level=h)
                elif s.startswith('Pregunta') or s.startswith('¿Pregunta'):
                    add_para(doc, s.replace('Pregunta', 'Pregunta:'), italic=True, space_before=Pt(6), space_after=Pt(1))
                elif s.startswith('Respuesta'):
                    add_para(doc, s.replace('Respuesta', 'Respuesta:'), indent=Cm(1.0), space_after=Pt(6))
                elif 'A LA GLORIA' in s or 'A la Gloria' in s:
                    add_para(doc, s, italic=True, font_size=Pt(10), alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=Pt(12))
                elif 'RITO ANTIGUO' in s:
                    add_para(doc, s, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, auto_bold=False, space_after=Pt(4))
                elif 'ORDEN DE LOS RITOS UNIDOS' in s:
                    add_para(doc, s, italic=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(4))
                elif s.startswith('GRAN LOGIA'):
                    add_para(doc, s, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, auto_bold=False, font_size=Pt(9))
                elif 'Ritos reunificados' in s or 'Ritos unificados' in s:
                    add_para(doc, s, italic=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, font_size=Pt(9), space_after=Pt(8))
                elif s.startswith('PROCLAMACIÓN') or s.startswith('PROCLAMACION') or s.startswith('CAPÍTULO') or s.startswith('CAPITULO') or s.startswith('CAP.'):
                    add_para(doc, s, bold=True, font_size=Pt(11), alignment=WD_ALIGN_PARAGRAPH.CENTER, auto_bold=False)
                elif s.startswith('TÍTULO') or s.startswith('TITULO') or s.startswith('DISPOSICIONES'):
                    add_para(doc, s, bold=True, auto_bold=False)
                elif s.startswith('CHEQUETET') or s.startswith('MAK-BENAH') or s.startswith('ARELICH'):
                    add_para(doc, s, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, auto_bold=False, space_before=Pt(8))
                elif s.startswith('Amen') or s.startswith('Soberano Maestro'):
                    add_para(doc, s, italic=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=Pt(6))
                elif re.match(r'^[\(\).\u2794\*].*[\)]$', s):
                    add_para(doc, s, italic=True, auto_bold=False)
                elif s.startswith('Art.') or s.startswith('Articulo') or s.startswith('ART.'):
                    add_para(doc, s, indent=Cm(0.5))
                else:
                    add_para(doc, s)


    return doc


def main():
    print("Cargando fuente completa...")
    source_text = load_source()
    
    print("Extrayendo paginas...")
    all_pages = extract_pages(source_text)
    print(f"  {len(all_pages)} paginas extraidas: {min(all_pages.keys()) if all_pages else '?'} - {max(all_pages.keys()) if all_pages else '?'}")

    os.makedirs(os.path.join(OUTPUT, 'doc'), exist_ok=True)
    
    completo_cfg = {
        'id': 'Rituales_Simbolicos_Completo',
        'title': 'Rituales e Instrucciones de los Grados Simbolicos - Obra Completa',
        'pages': list(range(2, 391)),
    }
    
    print(f'\nGenerando DOCX profesional...')
    doc = generate_docx(all_pages, completo_cfg)
    
    # Save to temp first to avoid permission issues
    temp_path = os.path.join(os.environ['TEMP'], 'ritual_completo.docx')
    doc.save(temp_path)
    
    target = os.path.join(OUTPUT, 'doc', 'Rituales_Simbolicos_Completo.docx')
    import shutil
    shutil.copy2(temp_path, target)
    os.remove(temp_path)
    print(f'  DOCX: Rituales_Simbolicos_Completo.docx ({os.path.getsize(target) / 1024:.0f} KB)')
    
    # Also generate standalone versions
    standalone_docs = [
        {'id': 'Ritual_1_Instruccion_1er_Grado_Aprendiz', 'title': 'Instruccion del 1er Grado - Aprendiz', 'pages': [68, 70, 71, 73, 74, 75]},
        {'id': 'Ritual_2_Instruccion_2o_Grado_Companero', 'title': 'Instruccion del 2o Grado - Companero', 'pages': [133, 134, 135, 136]},
        {'id': 'Ritual_3_Maestro', 'title': 'Ritual y Instruccion del 3er Grado - Maestro', 'pages': list(range(137, 178))},
        {'id': 'Grandes_Constituciones', 'title': 'Grandes Constituciones y Reglamentos Generales', 'pages': list(range(321, 391))},
    ]
    
    for cfg in standalone_docs:
        print(f'\n{cfg["title"]}')
        doc = generate_docx(all_pages, cfg)
        tp = os.path.join(os.environ['TEMP'], f'{cfg["id"]}.docx')
        doc.save(tp)
        tg = os.path.join(OUTPUT, 'doc', f'{cfg["id"]}.docx')
        import shutil
        shutil.copy2(tp, tg)
        os.remove(tp)
        print(f'  DOCX: {cfg["id"]}.docx')

    print('\n¡Todo generado!')


if __name__ == '__main__':
    main()
